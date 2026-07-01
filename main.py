"""
标书制作工具 - FastAPI REST API
确认驱动的工作流：Agent 做一步 → 用户确认/纠正 → 继续

所有业务逻辑委托给 Orchestrator。
"""
import json
import os
import re
import shutil
import time
import uuid
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Optional

import yaml
from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, PlainTextResponse
from pydantic import BaseModel
from sse_starlette.sse import EventSourceResponse

from models import init_db, get_session, Project, Tender, Material, MaterialUsage
from orchestrator import Orchestrator, WorkflowStep
from agents.bid_parser.pipeline import (
    BidLLMClient,
    BidParsePipeline,
    full_result_to_k01_k14,
)
from agents.bid_parser.pdf_extractor import get_file_info
from agents.bid_parser.marker_scanner import (
    extract_pages,
    scan_markers,
    summarize_markers,
)
from agents.bid_parser.schema import k_field_items_with_pages, k_field_value

# ---- 环境变量加载 ----
# 优先从 .env 读 TENDER_DEEPSEEK_API_KEY 等敏感信息（已在 .gitignore 中）
load_dotenv()

# ---- 配置加载 ----
_CONFIG_DIR = Path(__file__).parent
CONFIG_PATH = _CONFIG_DIR / "config.yaml"
if not CONFIG_PATH.exists():
    CONFIG_PATH = Path.home() / "tender-tool" / "config.yaml"
os.environ.setdefault("TENDER_CONFIG_PATH", str(CONFIG_PATH))


def load_config():
    with open(CONFIG_PATH, encoding="utf-8") as f:
        return yaml.safe_load(f)


config = load_config()
tender_config = config.get("tender_tool", {})

PROJECTS_DIR = Path(tender_config.get("projects_dir", "./projects"))
MATERIALS_DIR = Path(tender_config.get("materials_dir", "./materials"))
EXPORTS_DIR = Path(tender_config.get("exports_dir", "./exports"))

PROJECTS_DIR.mkdir(parents=True, exist_ok=True)
MATERIALS_DIR.mkdir(parents=True, exist_ok=True)
EXPORTS_DIR.mkdir(parents=True, exist_ok=True)

# ---- FastAPI App ----
app = FastAPI(title="标书制作工具", version="2.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

init_db()


# ===========================
# Pydantic 模型
# ===========================

class CreateProjectRequest(BaseModel):
    name: str
    tender_file_name: str


class TenderParseConfirmRequest(BaseModel):
    project_id: int
    corrections: Optional[dict] = None


class GenerateDraftRequest(BaseModel):
    project_id: int
    tender_type: str = "main"
    confirmed_chapters: Optional[dict] = None
    need_sub_bid: bool = False


class ExportRequest(BaseModel):
    format: str  # "markdown" / "word" / "pdf"


class ParsedDataPatchItem(BaseModel):
    field_path: str
    value: Any
    note: Optional[str] = None
    source: str = "user"


class ParsedDataPatchRequest(BaseModel):
    patches: list[ParsedDataPatchItem]


class PriceCalculationRequest(BaseModel):
    lowest_price: float
    main_price: float
    competitor_price: float
    low_price_ratio: Optional[float] = None
    highest_limit: Optional[float] = None
    price_score_max: Optional[float] = None


class MarkdownContentRequest(BaseModel):
    content: str


class MatchChapterUpdateRequest(BaseModel):
    chapter_id: str
    file_path: Optional[str] = None
    material_title: Optional[str] = None
    match_score: str = "高"
    reason: str = "用户手动选择"


CORRECTION_FIELD_ALIASES = {
    "项目名称": "K01_项目名称",
    "项目名": "K01_项目名称",
    "招标编号": "K02_招标编号",
    "项目编号": "K02_招标编号",
    "采购编号": "K02_招标编号",
    "预算": "K04_预算金额",
    "预算金额": "K04_预算金额",
    "控制价": "K04_预算金额",
    "最高限价": "K04_预算金额",
    "评分标准": "K07_评分标准",
    "技术要求": "K08_技术要求",
    "商务资质要求": "K09_商务资质要求",
    "资质要求": "K09_商务资质要求",
    "星标项": "K10_星标项",
    "废标条款": "K11_废标条款",
    "章节模板要求": "K12_章节模板要求",
    "模板要求": "K12_章节模板要求",
    "偏离表格式要求": "K13_偏离表格式要求",
    "偏离表要求": "K13_偏离表格式要求",
    "演示要求": "K14_演示要求",
}


class CreateMaterialRequest(BaseModel):
    title: str
    category: str
    description: str
    content: str
    content_type: str = "markdown"
    tags: Optional[str] = None


STANDARD_MATERIAL_CATEGORIES = [
    "01_公司资质",
    "02_业绩案例",
    "03_技术方案",
    "04_实施方案",
    "05_商务文件",
    "06_其他",
]


# ===========================
# 路由
# ===========================

@app.get("/")
def root():
    return {"message": "标书制作工具 API v3 (多 Agent)", "version": "2.0.0"}


# ---- 项目管理 ----

@app.post("/projects")
def create_project(req: CreateProjectRequest):
    safe_name = __import__("re").sub(r"[^\w\-]", "_", req.name)
    timestamp = datetime.now().strftime("%Y%m%d")
    project_dir = PROJECTS_DIR / f"{timestamp}_{safe_name}"
    project_dir.mkdir(parents=True, exist_ok=True)
    tender_path = project_dir / req.tender_file_name

    session = get_session()
    project = Project(
        name=req.name,
        tender_file_path=str(tender_path),
        status="parsing",
    )
    session.add(project)
    session.commit()
    session.refresh(project)

    return {
        "project_id": project.id,
        "project_name": project.name,
        "tender_file_path": str(tender_path),
        "message": f"项目已创建，请将招标文件上传到：{tender_path}"
    }


def _project_dir_for_delete(project: Project) -> Optional[Path]:
    """Return the project runtime directory only when it is safe to remove."""
    if not project.tender_file_path:
        return None

    project_dir = Path(project.tender_file_path).expanduser().parent.resolve()
    projects_root = PROJECTS_DIR.expanduser().resolve()
    try:
        project_dir.relative_to(projects_root)
    except ValueError:
        return None
    if project_dir == projects_root:
        return None
    return project_dir


@app.get("/projects")
def list_projects():
    session = get_session()
    projects = session.query(Project).order_by(Project.created_at.desc()).all()
    return [
        {
            "id": p.id,
            "name": p.name,
            "status": p.status,
            "created_at": p.created_at.isoformat(),
        }
        for p in projects
    ]


def _tender_summary(tender: Tender) -> dict:
    return {
        "id": tender.id,
        "type": tender.type,
        "status": tender.status,
        "draft_path": tender.draft_path,
        "created_at": tender.created_at.isoformat() if tender.created_at else None,
    }


def _active_main_tender(session, project_id: int) -> Tender | None:
    tenders = (
        session.query(Tender)
        .filter(Tender.project_id == project_id, Tender.type == "main")
        .order_by(Tender.id.desc())
        .all()
    )
    if not tenders:
        return None
    return next((t for t in tenders if t.draft_path), tenders[0])


TOOL_READINESS = {
    "parseTender": {
        "requires": ["tender_file"],
        "recoverable_action": "请先上传招标文件",
    },
    "outlineDesign": {
        "requires": ["parsed_data"],
        "recoverable_action": "请先解析招标文件",
    },
    "matchMaterials": {
        "requires": ["parsed_data", "outline"],
        "recoverable_action": "请先生成并确认提纲",
    },
    "generateTender": {
        "requires": ["outline", "matched_chapters"],
        "recoverable_action": "请先完成提纲设计和材料匹配",
    },
    "reviewTender": {
        "requires": ["draft"],
        "recoverable_action": "请先生成主标书",
    },
    "exportOutline": {
        "requires": ["outline"],
        "recoverable_action": "请先生成提纲",
    },
    "exportTender": {
        "requires": ["draft"],
        "recoverable_action": "请先生成主标书",
    },
}


def _project_parsed_data(project: Project) -> dict[str, Any]:
    if not project.parsed_data:
        return {}
    try:
        return json.loads(project.parsed_data)
    except json.JSONDecodeError:
        return {}


def _project_outline(parsed: dict[str, Any]) -> list[dict[str, Any]]:
    return (
        parsed.get("_confirmed_outline")
        or parsed.get("_generated_outline")
        or []
    )


def _check_tool_readiness(
    project: Project,
    tool_name: str,
    session=None,
) -> dict[str, Any]:
    spec = TOOL_READINESS.get(tool_name)
    if not spec:
        return {"ok": True, "tool": tool_name, "missing": []}

    parsed = _project_parsed_data(project)
    outline = _project_outline(parsed)
    missing: list[str] = []

    for requirement in spec["requires"]:
        if requirement == "tender_file":
            tender_path = Path(project.tender_file_path) if project.tender_file_path else None
            if not tender_path or not tender_path.exists():
                missing.append("tender_file")
        elif requirement == "parsed_data":
            if not parsed:
                missing.append("parsed_data")
        elif requirement == "outline":
            if not outline:
                missing.append("outline")
        elif requirement == "matched_chapters":
            if not parsed.get("chapters"):
                missing.append("matched_chapters")
        elif requirement == "draft":
            lookup_session = session or get_session()
            tender = _active_main_tender(lookup_session, project.id)
            draft_path = Path(tender.draft_path) if tender and tender.draft_path else None
            if not draft_path or not draft_path.exists():
                missing.append("draft")

    return {
        "ok": not missing,
        "tool": tool_name,
        "missing": missing,
        "requires": spec["requires"],
        "recoverable_action": spec["recoverable_action"],
        "message": (
            "当前项目数据不足，无法执行该工具。"
            if missing else "工具前置数据已满足。"
        ),
    }


def _normalize_export_format(fmt: str) -> str:
    value = (fmt or "word").lower().strip()
    aliases = {
        "docx": "word",
        "word": "word",
        "md": "markdown",
        "markdown": "markdown",
        "pdf": "pdf",
    }
    if value not in aliases:
        raise HTTPException(400, "format 仅支持 word/docx、markdown/md；PDF 后置")
    normalized = aliases[value]
    if normalized == "pdf":
        raise HTTPException(400, "PDF 导出暂不支持，请先导出 Word 或 Markdown")
    return normalized


def _export_format_from_message(msg: str) -> str:
    if re.search(r"pdf|PDF", msg):
        return "pdf"
    if re.search(r"word|docx|Word|WORD", msg):
        return "word"
    return "markdown"


def _parse_money_to_wan(value: Any) -> float | None:
    if value in (None, "", [], {}):
        return None
    if isinstance(value, dict):
        for key in ("amount", "value", "budget", "price"):
            if key in value:
                parsed = _parse_money_to_wan(value.get(key))
                if parsed is not None:
                    return parsed
        return None
    text = str(value).replace(",", "").strip()
    m = re.search(r"(\d+(?:\.\d+)?)\s*(万元|万|元)?", text)
    if not m:
        return None
    amount = float(m.group(1))
    unit = m.group(2) or ""
    return amount / 10000 if unit == "元" else amount


def _format_wan(value: float | None) -> str:
    if value is None:
        return "—"
    return f"{value:,.2f} 万元" if abs(value) >= 100 else f"{value:.4f} 万元"


def _extract_price_score_max(parsed: dict[str, Any]) -> float | None:
    scoring = parsed.get("scoring") or {}
    if isinstance(scoring, dict):
        for dim in scoring.get("dimensions") or []:
            if not isinstance(dim, dict):
                continue
            name = str(dim.get("name") or "")
            if any(kw in name for kw in ("价格", "报价", "投标报价")):
                score = dim.get("max_score")
                if isinstance(score, (int, float)) and score > 0:
                    return float(score)
                parsed_score = _parse_money_to_wan(score)
                if parsed_score:
                    return parsed_score
        ratio = scoring.get("price_ratio")
        if isinstance(ratio, (int, float)) and ratio > 0:
            return float(ratio)

    text = str(k_field_value(parsed.get("K07_评分标准")) or "")
    m = re.search(r"(?:价格|报价)[^。；;\n]{0,20}?(\d+(?:\.\d+)?)\s*分", text)
    return float(m.group(1)) if m else None


def _extract_low_price_ratio(parsed: dict[str, Any], msg: str = "") -> float | None:
    text = msg or ""
    scoring = parsed.get("scoring") or {}
    if isinstance(scoring, dict):
        low_price = scoring.get("low_price_review") or {}
        if isinstance(low_price, dict):
            text += " " + str(low_price.get("trigger") or "")
    for field_key in ("K07_评分标准", "K10_星标项", "K11_废标条款"):
        value = k_field_value(parsed.get(field_key))
        if isinstance(value, list):
            text += " " + " ".join(str(v) for v in value)
        elif value:
            text += " " + str(value)

    patterns = [
        r"(?:低价|异常低价|低价审核|低价线|触发比例|预警线|低价阈值)[^%]{0,40}?(\d+(?:\.\d+)?)\s*%",
        r"(?:最高限价|预算|控制价)[^%]{0,40}?(\d+(?:\.\d+)?)\s*%",
    ]
    for pattern in patterns:
        m = re.search(pattern, text)
        if m:
            value = float(m.group(1))
            return value / 100 if value > 1 else value
    return None


def _extract_labeled_price(msg: str, labels: list[str]) -> float | None:
    label_pattern = "|".join(re.escape(label) for label in labels)
    m = re.search(
        rf"(?:{label_pattern})[^\d]{{0,20}}(\d+(?:\.\d+)?)\s*(万元|万|元)?",
        msg,
    )
    if not m:
        return None
    return _parse_money_to_wan("".join(part or "" for part in m.groups()))


def _extract_price_calc_inputs(msg: str, parsed: dict[str, Any]) -> dict[str, Any]:
    inputs: dict[str, Any] = {
        "lowest_price": _extract_labeled_price(msg, ["最低报价", "最低价", "最低", "基准价"]),
        "main_price": _extract_labeled_price(msg, ["主标报价", "主标", "我方报价", "我方", "本方报价"]),
        "competitor_price": _extract_labeled_price(msg, ["竞争对手报价", "对手报价", "竞争对手", "对手"]),
        "low_price_ratio": _extract_low_price_ratio(parsed, msg),
        "highest_limit": None,
        "price_score_max": _extract_price_score_max(parsed),
    }
    explicit_limit = re.search(
        r"(?:最高限价|预算|控制价)[^\d]{0,12}(\d+(?:\.\d+)?)\s*(万元|万|元)?",
        msg,
    )
    if explicit_limit:
        inputs["highest_limit"] = _parse_money_to_wan(
            "".join(part or "" for part in explicit_limit.groups())
        )
    else:
        base = parsed.get("base") or {}
        inputs["highest_limit"] = (
            _parse_money_to_wan(k_field_value(parsed.get("K04_预算金额")))
            or _parse_money_to_wan(base.get("budget") if isinstance(base, dict) else None)
        )

    m_score = re.search(r"(?:价格分|报价分|价格满分|价格权重)[^\d]{0,12}(\d+(?:\.\d+)?)\s*(?:分|%)?", msg)
    if m_score:
        inputs["price_score_max"] = float(m_score.group(1))

    m_ratio = re.search(r"(?:低价比例|异常低价比例|低价线|低价阈值|触发比例)[^\d]{0,12}(\d+(?:\.\d+)?)\s*%", msg)
    if m_ratio:
        inputs["low_price_ratio"] = float(m_ratio.group(1)) / 100

    money_mentions = [
        _parse_money_to_wan("".join(part or "" for part in m.groups()))
        for m in re.finditer(r"(\d+(?:\.\d+)?)\s*(万元|万|元)", msg)
    ]
    if any(inputs[key] is None for key in ("lowest_price", "main_price", "competitor_price")) and len(money_mentions) >= 3:
        quote_values = [v for v in money_mentions if v is not None][-3:]
        if len(quote_values) == 3:
            inputs["lowest_price"] = inputs["lowest_price"] or quote_values[0]
            inputs["main_price"] = inputs["main_price"] or quote_values[1]
            inputs["competitor_price"] = inputs["competitor_price"] or quote_values[2]
    return inputs


def _calculate_price_scores(
    parsed: dict[str, Any],
    *,
    lowest_price: float | None,
    main_price: float | None,
    competitor_price: float | None,
    low_price_ratio: float | None = None,
    highest_limit: float | None = None,
    price_score_max: float | None = None,
) -> dict[str, Any]:
    price_score_max = price_score_max or _extract_price_score_max(parsed)
    low_price_ratio = low_price_ratio or _extract_low_price_ratio(parsed)
    highest_limit = highest_limit or _parse_money_to_wan(k_field_value(parsed.get("K04_预算金额")))

    missing = []
    for key, value, label in [
        ("lowest_price", lowest_price, "最低报价"),
        ("main_price", main_price, "主标报价"),
        ("competitor_price", competitor_price, "竞争对手报价"),
    ]:
        if value is None or value <= 0:
            missing.append({"field": key, "label": label})
    if highest_limit is None or highest_limit <= 0:
        missing.append({"field": "highest_limit", "label": "最高限价"})
    if price_score_max is None or price_score_max <= 0:
        missing.append({"field": "price_score_max", "label": "价格分满分"})
    if low_price_ratio is None or low_price_ratio <= 0:
        missing.append({"field": "low_price_ratio", "label": "异常低价触发比例"})
    if missing:
        return {
            "ok": False,
            "missing": missing,
            "message": "价格测算缺少必要参数",
            "action_hint": "请补充：最低报价、主标报价、竞争对手报价、最高限价、价格分满分或低价比例。",
        }

    benchmark = min(lowest_price, main_price, competitor_price)
    low_price_threshold = highest_limit * low_price_ratio
    rows = []
    for key, label, price in [
        ("lowest", "最低报价", lowest_price),
        ("main", "主标报价", main_price),
        ("competitor", "竞争对手报价", competitor_price),
    ]:
        score = price_score_max * benchmark / price
        rows.append({
            "key": key,
            "label": label,
            "price": round(price, 4),
            "price_display": _format_wan(price),
            "score": round(score, 4),
            "score_display": f"{score:.4f}",
            "score_diff_to_main": None,
            "price_diff_to_main": None,
            "triggers_low_price": price < low_price_threshold,
        })

    main_row = next(row for row in rows if row["key"] == "main")
    for row in rows:
        row["score_diff_to_main"] = round(row["score"] - main_row["score"], 4)
        row["price_diff_to_main"] = round(row["price"] - main_row["price"], 4)

    score_by_key = {row["key"]: row["score"] for row in rows}
    return {
        "ok": True,
        "method": "最低报价基准法",
        "formula": "价格分 = 价格分满分 × 最低报价 / 当前报价",
        "price_score_max": round(price_score_max, 4),
        "benchmark_price": round(benchmark, 4),
        "benchmark_price_display": _format_wan(benchmark),
        "highest_limit": round(highest_limit, 4),
        "highest_limit_display": _format_wan(highest_limit),
        "low_price_ratio": round(low_price_ratio, 6),
        "low_price_ratio_display": f"{low_price_ratio * 100:.2f}%",
        "low_price_threshold": round(low_price_threshold, 4),
        "low_price_threshold_display": _format_wan(low_price_threshold),
        "rows": rows,
        "pairwise_score_diff": {
            "main_vs_lowest": round(score_by_key["main"] - score_by_key["lowest"], 4),
            "main_vs_competitor": round(score_by_key["main"] - score_by_key["competitor"], 4),
            "competitor_vs_lowest": round(score_by_key["competitor"] - score_by_key["lowest"], 4),
        },
        "any_low_price_risk": any(row["triggers_low_price"] for row in rows),
        "message": "价格测算完成",
    }


def _price_calculator_defaults(parsed: dict[str, Any]) -> dict[str, Any]:
    base = parsed.get("base") or {}
    highest_limit = (
        _parse_money_to_wan(k_field_value(parsed.get("K04_预算金额")))
        or _parse_money_to_wan(base.get("budget") if isinstance(base, dict) else None)
    )
    low_price_ratio = _extract_low_price_ratio(parsed)
    price_score_max = _extract_price_score_max(parsed)
    missing = []
    if highest_limit is None or highest_limit <= 0:
        missing.append({"field": "highest_limit", "label": "最高限价"})
    if low_price_ratio is None or low_price_ratio <= 0:
        missing.append({"field": "low_price_ratio", "label": "异常低价触发比例"})
    if price_score_max is None or price_score_max <= 0:
        missing.append({"field": "price_score_max", "label": "价格分满分"})
    return {
        "highest_limit": highest_limit,
        "highest_limit_display": _format_wan(highest_limit),
        "low_price_ratio": low_price_ratio,
        "low_price_ratio_display": f"{low_price_ratio * 100:.2f}%" if low_price_ratio else "—",
        "price_score_max": price_score_max,
        "price_score_max_display": f"{price_score_max:.2f} 分" if price_score_max else "—",
        "missing": missing,
        "ok": not missing,
    }


def _should_run_price_calculator(msg: str) -> bool:
    return bool(re.search(r"价格.*(测算|计算|打分|得分)|报价.*(测算|计算|打分|得分)|异常低价.*(测算|计算|判断)|低价线", msg))


def _should_start_parse_from_message(msg: str, project_status: str | None = None) -> bool:
    """Return True when a chat message is asking to parse the uploaded tender file."""
    text = (msg or "").strip()
    if not text:
        return False
    if re.search(r"重新解析|重跑解析|重新分析招标", text, re.IGNORECASE):
        return True
    if project_status != "parsing":
        return False
    if re.search(r"放好了|开始解析|解析招标|识别招标|分析招标|读取招标", text, re.IGNORECASE):
        return True
    mentions_file = re.search(r"上传|文件|招标文件|标书|pdf|docx|附件", text, re.IGNORECASE)
    asks_action = re.search(r"好了|完成|完了|开始|处理|看看|看一下|分析|识别|读取|继续", text)
    return bool(mentions_file and asks_action)


def _markdown_to_docx(markdown: str, output_path: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise HTTPException(500, "python-docx 未安装，无法导出 Word")

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line).strip(), style="List Number")
        elif line.strip() == "---":
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
        elif line.startswith(">"):
            doc.add_paragraph(line.lstrip("> ").strip(), style="Intense Quote")
        else:
            doc.add_paragraph(line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def _export_tender_file(tender: Tender, fmt: str) -> dict:
    export_format = _normalize_export_format(fmt)
    if not tender.draft_path:
        raise HTTPException(400, "标书尚未生成 draft.md，无法导出")

    draft_path = Path(tender.draft_path)
    if not draft_path.exists():
        raise HTTPException(404, f"draft.md 不存在: {draft_path}")

    markdown = draft_path.read_text(encoding="utf-8")
    if export_format == "markdown":
        export_path = EXPORTS_DIR / f"tender_{tender.id}.md"
        export_path.parent.mkdir(parents=True, exist_ok=True)
        export_path.write_text(markdown, encoding="utf-8")
        download_format = "markdown"
    else:
        export_path = EXPORTS_DIR / f"tender_{tender.id}.docx"
        _markdown_to_docx(markdown, export_path)
        download_format = "word"

    return {
        "format": export_format,
        "export_path": str(export_path),
        "download_url": f"/api/downloads/{tender.id}/{download_format}",
    }


def _volume_label(volume: str) -> str:
    return {
        "commercial": "商务文件",
        "technical": "技术文件",
    }.get(volume or "commercial", "商务文件")


def _outline_volume(volume: str) -> str:
    return "technical" if volume == "technical" else "commercial"


def _format_outline_refs(refs: list[dict[str, Any]] | None) -> str:
    if not refs:
        return "无"
    parts = []
    for ref in refs[:3]:
        page = ref.get("page")
        display_label = (ref.get("label") or "").strip()
        quote = (ref.get("quote") or "").strip()
        if page:
            label = f"P.{page}"
            if display_label:
                label += f"：{display_label[:80]}"
            elif quote:
                label += f"：{quote[:80]}"
        elif display_label:
            label = display_label[:80]
        elif quote:
            label = quote[:80]
        else:
            label = "来源未定位"
        parts.append(label)
    return "；".join(parts)


def _format_scoring_title(name: Any, score: Any = None) -> str:
    title = re.sub(r"[；;:\s]+$", "", str(name or "").strip())
    if not title:
        return ""
    if score in (None, ""):
        return title
    score_text = re.sub(r"\s+", "", str(score).strip())
    if not score_text:
        return title
    if re.search(r"(分|%)$", score_text):
        return f"{title} {score_text}"
    return f"{title} {score_text}分"


def _display_scoring_label(value: Any) -> str:
    text = str(value or "").strip()
    text = re.sub(r"[；;]\s*(\d+(?:\.\d+)?)\s*$", r" \1分", text)
    text = re.sub(r"\s+(\d+(?:\.\d+)?)\s*$", r" \1分", text)
    return text.replace("分分", "分")


def _is_price_scoring_text(value: Any) -> bool:
    text = str(value or "")
    return any(kw in text for kw in ("报价", "价格", "投标报价", "开标一览", "分项报价"))


def _outline_chapter_text(chapter: dict[str, Any]) -> str:
    return " ".join(
        [str(chapter.get("title") or "")]
        + [
            str(sub.get("title") or "")
            for sub in chapter.get("subsections") or []
            if isinstance(sub, dict)
        ]
    )


def _price_scoring_titles(parsed: dict[str, Any]) -> list[str]:
    scoring = parsed.get("scoring") or {}
    titles: list[str] = []
    for dim in scoring.get("dimensions") or []:
        if not isinstance(dim, dict):
            continue
        dim_name = str(dim.get("name") or "")
        dim_score = dim.get("max_score")
        dim_titles: list[str] = []
        for sub in dim.get("sub_items") or []:
            if not isinstance(sub, dict):
                continue
            sub_name = str(sub.get("name") or "")
            if _is_price_scoring_text(f"{dim_name} {sub_name}"):
                label = _format_scoring_title(sub_name or dim_name, sub.get("score") or dim_score)
                if label and label not in dim_titles:
                    dim_titles.append(label)
        if dim_titles:
            for label in dim_titles:
                if label not in titles:
                    titles.append(label)
        elif _is_price_scoring_text(dim_name):
            label = _format_scoring_title(dim_name, dim_score)
            if label and label not in titles:
                titles.append(label)
    return titles


def _normalize_outline_for_display(
    outline: list[dict[str, Any]],
    parsed: dict[str, Any] | None = None,
) -> list[dict[str, Any]]:
    normalized: list[dict[str, Any]] = []
    for idx, source in enumerate(outline, 1):
        if not isinstance(source, dict):
            continue
        chapter = dict(source)
        chapter["no"] = idx
        if chapter.get("title"):
            chapter["title"] = _display_scoring_label(chapter["title"])
        chapter["subsections"] = [
            {
                **sub,
                "title": _display_scoring_label(sub.get("title")),
            }
            if isinstance(sub, dict)
            else {"id": f"{chapter.get('id') or f'ch{idx}'}.{j}", "title": _display_scoring_label(sub)}
            for j, sub in enumerate(chapter.get("subsections") or [], 1)
        ]

        chapter_text = _outline_chapter_text(chapter)
        filtered_refs = []
        for ref in chapter.get("scoring_refs") or []:
            if not isinstance(ref, dict):
                continue
            ref_copy = dict(ref)
            if ref_copy.get("label"):
                ref_copy["label"] = _display_scoring_label(ref_copy["label"])
            ref_text = f"{ref_copy.get('label') or ''} {ref_copy.get('quote') or ''}"
            if _is_price_scoring_text(ref_text) and not _is_price_scoring_text(chapter_text):
                continue
            filtered_refs.append(ref_copy)
        chapter["scoring_refs"] = filtered_refs
        normalized.append(chapter)

    if parsed:
        price_titles = _price_scoring_titles(parsed)
        if price_titles:
            target = next(
                (ch for ch in normalized if _is_price_scoring_text(_outline_chapter_text(ch))),
                None,
            )
            if target is None:
                target = {
                    "id": f"ch{len(normalized) + 1}",
                    "no": len(normalized) + 1,
                    "title": "投标报价文件",
                    "volume": "commercial",
                    "category": "05_商务文件",
                    "subsections": [],
                    "source": "scoring",
                    "requirement_refs": [],
                    "scoring_refs": [],
                }
                normalized.append(target)
            subs = target.setdefault("subsections", [])
            existing_titles = {
                str(sub.get("title") or "")
                for sub in subs
                if isinstance(sub, dict)
            }
            for title in price_titles:
                if title not in existing_titles:
                    subs.append({
                        "id": f"{target.get('id') or 'ch'}.{len(subs) + 1}",
                        "title": title,
                        "source": "scoring",
                    })
                    existing_titles.add(title)

    for idx, chapter in enumerate(normalized, 1):
        chapter["no"] = idx
        chapter["id"] = chapter.get("id") or f"ch{idx}"
        for sub_idx, sub in enumerate(chapter.get("subsections") or [], 1):
            if isinstance(sub, dict):
                sub["id"] = sub.get("id") or f"{chapter['id']}.{sub_idx}"
    return normalized


def _normalize_project_parsed_data_for_display(parsed: dict[str, Any]) -> dict[str, Any]:
    out = dict(parsed)
    for key in ("_generated_outline", "_confirmed_outline"):
        value = out.get(key)
        if isinstance(value, list):
            out[key] = _normalize_outline_for_display(value, out)
    return out


def _build_outline_markdown(project: Project, outline: list[dict[str, Any]]) -> str:
    parsed = json.loads(project.parsed_data) if project.parsed_data else {}
    outline = _normalize_outline_for_display(outline, parsed)
    lines = [
        f"# {project.project_name or project.name} - 投标文件提纲",
        "",
        "<!-- 编辑说明：直接修改章节标题、分类和小节；保存后工作流会使用此提纲。 -->",
        "",
        f"- 项目 ID: {project.id}",
    ]
    if project.tender_no:
        lines.append(f"- 招标编号: {project.tender_no}")
    lines.extend([
        "",
    ])

    grouped = {"commercial": [], "technical": []}
    for ch in outline:
        volume = _outline_volume(ch.get("volume", "commercial"))
        grouped[volume].append(ch)

    for volume, chapters in grouped.items():
        if not chapters:
            continue
        lines.extend(["", f"## {_volume_label(volume)}", ""])
        for no, ch in enumerate(chapters, 1):
            title = ch.get("title") or f"第{no}章"
            category = ch.get("category") or "未分类"
            source = ch.get("source") or "unknown"
            lines.append(f"### {title}")
            lines.append(f"- 分类: `{category}`")
            lines.append(f"- 来源: `{source}`")
            lines.append(f"- 投标要求位置: {_format_outline_refs(ch.get('requirement_refs'))}")
            lines.append(f"- 评分位置: {_format_outline_refs(ch.get('scoring_refs'))}")
            for sub in ch.get("subsections") or []:
                sub_title = sub.get("title") if isinstance(sub, dict) else str(sub)
                if sub_title:
                    lines.append(f"- {sub_title}")
            lines.append("")

    return "\n".join(lines).strip() + "\n"


def _parse_outline_markdown(markdown: str, existing: list[dict[str, Any]] | None = None) -> list[dict[str, Any]]:
    existing_by_title = {
        str(ch.get("title") or "").strip(): ch
        for ch in (existing or [])
        if str(ch.get("title") or "").strip()
    }
    outline: list[dict[str, Any]] = []
    volume = "commercial"
    current: dict[str, Any] | None = None

    def push_current():
        nonlocal current
        if current and str(current.get("title") or "").strip():
            outline.append(current)
        current = None

    for raw in markdown.splitlines():
        line = raw.strip()
        if not line or line.startswith("<!--"):
            continue
        if re.match(r"^#{1,2}\s+.*商务文件", line):
            push_current()
            volume = "commercial"
            continue
        if re.match(r"^#{1,2}\s+.*技术文件", line):
            push_current()
            volume = "technical"
            continue
        heading = re.match(r"^###\s+(.+)$", line)
        numbered = re.match(r"^\d+[\.\、]\s+\*\*(.+?)\*\*", line)
        if heading or numbered:
            push_current()
            title = (heading.group(1) if heading else numbered.group(1)).strip()
            prior = existing_by_title.get(title, {})
            current = {
                "id": prior.get("id") or f"ch{len(outline) + 1}",
                "no": len(outline) + 1,
                "title": title,
                "volume": prior.get("volume") or volume,
                "category": prior.get("category") or ("03_技术方案" if volume == "technical" else "05_商务文件"),
                "subsections": [],
                "source": prior.get("source") or "user_edited",
            }
            for key in ("requirement_refs", "scoring_refs"):
                if key in prior:
                    current[key] = prior[key]
            continue
        if not current:
            continue
        category = re.match(r"^-\s*分类[:：]\s*`?([^`]+?)`?\s*$", line)
        if category:
            value = category.group(1).strip()
            if value in STANDARD_MATERIAL_CATEGORIES:
                current["category"] = value
            continue
        source = re.match(r"^-\s*来源[:：]\s*`?([^`]+?)`?\s*$", line)
        if source:
            current["source"] = source.group(1).strip() or "user_edited"
            continue
        if re.match(r"^-\s*(投标要求位置|评分位置)[:：]", line):
            continue
        bullet = re.match(r"^-\s+(.+)$", line)
        if bullet:
            title = bullet.group(1).strip()
            if title:
                current.setdefault("subsections", []).append({
                    "id": f"{current.get('id')}.{len(current.get('subsections') or []) + 1}",
                    "title": title,
                })
    push_current()
    for idx, ch in enumerate(outline, 1):
        ch["no"] = idx
        if not ch.get("id"):
            ch["id"] = f"ch{idx}"
    return outline


def _deviation_items_from_parsed(parsed: dict[str, Any]) -> dict[str, list[str]]:
    confirmed = parsed.get("_confirmed_deviation_items")
    pending = parsed.get("_pending_deviation_items")
    if isinstance(pending, dict):
        source = pending
    elif isinstance(confirmed, dict):
        source = confirmed
    else:
        source = {
            "business": Orchestrator._business_deviation_items(parsed),
            "technical": Orchestrator._technical_deviation_items(parsed),
        }
        star_items = [
            item for item, _page in k_field_items_with_pages(parsed.get("K10_星标项"))
        ]
        for item in star_items:
            target = Orchestrator._deviation_item_target(item)
            if target in source and item not in source[target]:
                source[target].append(item)
    return {
        "business": [str(v).strip() for v in source.get("business", []) if str(v).strip()],
        "technical": [str(v).strip() for v in source.get("technical", []) if str(v).strip()],
    }


def _build_deviation_editor_markdown(project: Project, parsed: dict[str, Any]) -> str:
    items = _deviation_items_from_parsed(parsed)
    project_name = k_field_value(parsed.get("K01_项目名称")) or project.project_name or project.name
    lines = [
        f"# {project_name} - 商务/技术条款偏离表",
        "",
        "<!-- 编辑说明：修改表格第二列的招标条款/要求；保存后确认生成会使用这些条目。 -->",
        "",
        "## 商务条款偏离表",
        "",
        "| 序号 | 招标条款/要求 | 投标响应 | 偏离情况 | 说明 |",
        "| --- | --- | --- | --- | --- |",
    ]
    for idx, item in enumerate(items["business"], 1):
        lines.append(f"| {idx} | {Orchestrator._table_cell(item)} | 完全响应 | 无偏离 | 按招标文件及商务资质材料响应 |")
    if not items["business"]:
        lines.append("| 1 | 待补充商务条款 | 待补充 | 待确认 | 保存前请替换本行 |")
    lines.extend([
        "",
        "## 技术条款偏离表",
        "",
        "| 序号 | 招标技术要求 | 投标响应 | 偏离情况 | 说明 |",
        "| --- | --- | --- | --- | --- |",
    ])
    for idx, item in enumerate(items["technical"], 1):
        lines.append(f"| {idx} | {Orchestrator._table_cell(item)} | 完全响应 | 无偏离 | 技术方案章节已响应 |")
    if not items["technical"]:
        lines.append("| 1 | 待补充技术条款 | 待补充 | 待确认 | 保存前请替换本行 |")
    return "\n".join(lines).strip() + "\n"


def _parse_deviation_markdown(markdown: str) -> dict[str, list[str]]:
    target: str | None = None
    result = {"business": [], "technical": []}
    for raw in markdown.splitlines():
        line = raw.strip()
        if re.match(r"^##\s+商务", line):
            target = "business"
            continue
        if re.match(r"^##\s+技术", line):
            target = "technical"
            continue
        if not target or not line.startswith("|") or "---" in line:
            continue
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if cells and cells[0] == "序号":
            continue
        if len(cells) >= 2:
            item = cells[1].replace("\\|", "|").strip()
            if item and not item.startswith("待补充") and item not in result[target]:
                result[target].append(item)
    return result


def _export_outline_file(project: Project, fmt: str = "markdown") -> dict:
    export_format = _normalize_export_format(fmt)
    if export_format != "markdown":
        raise HTTPException(400, "提纲阶段暂只支持 Markdown 导出")
    parsed = json.loads(project.parsed_data) if project.parsed_data else {}
    outline = (
        parsed.get("_confirmed_outline")
        or parsed.get("_generated_outline")
        or []
    )
    if not outline:
        raise HTTPException(400, "尚未生成提纲，无法导出")
    markdown = _build_outline_markdown(project, outline)
    export_path = EXPORTS_DIR / f"outline_{project.id}.md"
    export_path.parent.mkdir(parents=True, exist_ok=True)
    export_path.write_text(markdown, encoding="utf-8")
    return {
        "format": "markdown",
        "export_path": str(export_path),
        "download_url": f"/api/downloads/outlines/{project.id}/markdown",
    }


@app.get("/projects/{project_id}")
def get_project(project_id: int):
    session = get_session()
    project = session.get(Project, project_id)
    if not project:
        raise HTTPException(404, "项目不存在")

    tenders = (
        session.query(Tender)
        .filter(Tender.project_id == project_id)
        .order_by(Tender.id.desc())
        .all()
    )
    active_main_tender = next(
        (t for t in tenders if t.type == "main" and t.draft_path),
        next((t for t in tenders if t.type == "main"), None),
    )

    parsed_data = json.loads(project.parsed_data) if project.parsed_data else None
    if isinstance(parsed_data, dict):
        parsed_data = _normalize_project_parsed_data_for_display(parsed_data)

    return {
        "id": project.id,
        "name": project.name,
        "status": project.status,
        "tender_file_path": project.tender_file_path,
        "project_name": project.project_name,
        "tender_no": project.tender_no,
        "budget": project.budget,
        "deadline": project.deadline.isoformat() if project.deadline else None,
        "open_time": project.open_time.isoformat() if project.open_time else None,
        "tender_id": active_main_tender.id if active_main_tender else None,
        "active_main_tender_id": active_main_tender.id if active_main_tender else None,
        "tenders": [_tender_summary(t) for t in tenders],
        # 包含已解析数据（前端 ChatView 重新进入项目时恢复解析报告）
        "parsed_data": parsed_data,
        # 对话历史（UIMessage[]，useChat 初始化时回填）
        "messages": json.loads(project.messages_json) if project.messages_json else [],
    }


@app.get("/projects/{project_id}/outline/markdown")
def get_outline_markdown(project_id: int):
    session = get_session()
    project = session.get(Project, project_id)
    if not project:
        raise HTTPException(404, "项目不存在")
    parsed = json.loads(project.parsed_data) if project.parsed_data else {}
    outline = parsed.get("_confirmed_outline") or parsed.get("_generated_outline") or []
    if not outline:
        raise HTTPException(404, "尚未生成提纲")
    return PlainTextResponse(
        _build_outline_markdown(project, outline),
        media_type="text/markdown; charset=utf-8",
    )


@app.put("/projects/{project_id}/outline/markdown")
def save_outline_markdown(project_id: int, req: MarkdownContentRequest):
    session = get_session()
    project = session.get(Project, project_id)
    if not project:
        raise HTTPException(404, "项目不存在")
    parsed = json.loads(project.parsed_data) if project.parsed_data else {}
    existing = parsed.get("_confirmed_outline") or parsed.get("_generated_outline") or []
    outline = _parse_outline_markdown(req.content, existing)
    if not outline:
        raise HTTPException(400, "Markdown 中未解析到有效章节")
    parsed["_generated_outline"] = outline
    if project.status not in {"parsed", "outline_generating"}:
        parsed["_confirmed_outline"] = outline
    else:
        parsed.pop("_confirmed_outline", None)
        project.status = "outline_generating"
    project.parsed_data = json.dumps(parsed, ensure_ascii=False)
    session.commit()
    return {"message": "提纲已保存", "outline": outline, "total": len(outline)}


@app.get("/projects/{project_id}/deviation/markdown")
def get_deviation_markdown(project_id: int):
    session = get_session()
    project = session.get(Project, project_id)
    if not project:
        raise HTTPException(404, "项目不存在")
    parsed = json.loads(project.parsed_data) if project.parsed_data else {}
    if not parsed:
        raise HTTPException(404, "尚未解析招标文件")
    return PlainTextResponse(
        _build_deviation_editor_markdown(project, parsed),
        media_type="text/markdown; charset=utf-8",
    )


@app.put("/projects/{project_id}/deviation/markdown")
def save_deviation_markdown(project_id: int, req: MarkdownContentRequest):
    session = get_session()
    project = session.get(Project, project_id)
    if not project:
        raise HTTPException(404, "项目不存在")
    parsed = json.loads(project.parsed_data) if project.parsed_data else {}
    items = _parse_deviation_markdown(req.content)
    parsed["_pending_deviation_items"] = items
    parsed.pop("_confirmed_deviation_items", None)
    if project.status not in {"generating", "generated", "reviewing", "reviewed", "done"}:
        project.status = "deviation_preparing"
    project.parsed_data = json.dumps(parsed, ensure_ascii=False)
    session.commit()
    return {"message": "偏离表条目已保存", **items}


@app.delete("/projects/{project_id}")
def delete_project(project_id: int):
    session = get_session()
    project = session.get(Project, project_id)
    if not project:
        raise HTTPException(404, "项目不存在")

    project_dir = _project_dir_for_delete(project)

    try:
        # 清理关联数据：材料使用记录 → 标书 → 项目
        session.query(MaterialUsage).filter(MaterialUsage.project_id == project_id).delete()
        session.query(Tender).filter(Tender.project_id == project_id).delete()
        session.delete(project)
        if project_dir and project_dir.exists():
            shutil.rmtree(project_dir)
        session.commit()
    except OSError as exc:
        session.rollback()
        raise HTTPException(500, f"项目文件夹删除失败：{exc}") from exc
    except Exception:
        session.rollback()
        raise
    return {
        "message": "项目已删除",
        "deleted_project_dir": str(project_dir) if project_dir else None,
    }


@app.put("/projects/{project_id}/messages")
def save_messages(project_id: int, payload: dict):
    """保存对话历史（useChat.onFinish 回调）。"""
    session = get_session()
    project = session.get(Project, project_id)
    if not project:
        raise HTTPException(404, "项目不存在")
    messages = payload.get("messages", [])
    if not isinstance(messages, list):
        raise HTTPException(400, "messages 必须为数组")
    project.messages_json = json.dumps(messages, ensure_ascii=False)
    session.commit()
    return {"message": "已保存", "count": len(messages)}


# ---- 招标文件解析 ----

@app.post("/projects/{project_id}/upload")
async def upload_tender(project_id: int, file: UploadFile = File(...)):
    """上传招标文件（multipart/form-data）。

    保存到 project.tender_file_path 并将 status 置回 "parsing"。
    不自动触发解析——由用户说「放好了」后由 /parse 端点执行，
    保持"AI 做一步 → 用户确认"的节奏。

    实际保存路径会按上传文件后缀调整（避免 .pdf 后缀装 DOCX 内容的问题）。
    """
    session = get_session()
    project = session.get(Project, project_id)
    if not project:
        raise HTTPException(404, "项目不存在")

    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in {".pdf", ".docx"}:
        raise HTTPException(400, f"不支持的文件类型：{suffix}（仅接受 .pdf / .docx）")

    # 清洗文件名：去除路径分隔符和危险字符，保留中文/英文/数字/_-
    import re as _re
    raw_name = Path(file.filename or "tender").name
    safe_name = _re.sub(r"[^\w一-鿿\-.]", "_", raw_name)
    if not safe_name or safe_name.startswith("."):
        safe_name = f"tender{suffix}"

    content = await file.read()
    # 用原文件名（清洗后）覆盖占位 tender.pdf — 前端能显示实际名字
    base = Path(project.tender_file_path).parent / safe_name
    base.parent.mkdir(parents=True, exist_ok=True)
    base.write_bytes(content)
    # 同步更新 DB 中的路径，保证后续 /parse 走正确的解析分支
    project.tender_file_path = str(base)
    project.status = "parsing"
    session.commit()

    return {
        "message": "上传成功，可以直接说“开始解析招标文件”或描述“文件已上传，请帮我分析”。",
        "file_path": str(base),
        "file_size": len(content),
        "filename": safe_name,
    }


@app.post("/projects/{project_id}/parse")
def parse_tender(project_id: int, mode: Optional[str] = None):
    """由 ParserAgent 解析招标文件，提取 K01-K14 + 结构化数据。

    Query:
        mode: 解析模式覆盖 — auto/quick/full/manual。
              缺省走 config.yaml 的 parser.mode（默认 full）。
    """
    session = get_session()
    project = session.get(Project, project_id)
    if not project:
        raise HTTPException(404, "项目不存在")

    tender_path = Path(project.tender_file_path)
    if not tender_path.exists():
        raise HTTPException(400, f"招标文件不存在：{tender_path}")

    orch = Orchestrator(tender_config)
    orch.ctx.project_id = project_id
    orch.ctx.parsed_data["K01_项目名称"] = project.name
    if mode:
        orch.ctx.parser_mode_override = mode

    parsed = orch.agents["parser"].execute(orch.ctx)

    project.parsed_data = json.dumps(parsed, ensure_ascii=False)
    project.status = "parsed"
    session.commit()

    return {
        "message": "解析完成，请确认以下信息是否正确：",
        "parsed_data": parsed,
        "correction_hint": "如有错误，请告知我需要修改的内容"
    }


# ============================================================
# 分步解析：每步独立端点，前端可串行调用以显示进度
# ============================================================

# 内存缓存：project_id → 解析中间态
# 解析完成或项目删除时清除。重启服务会丢失（用户重新发起即可）
_parse_state: dict[int, dict] = {}


def _make_pipeline() -> BidParsePipeline:
    """构造一个独立的 pipeline 实例（无状态，可随时新建）。"""
    ai = tender_config.get("ai", {})
    parser_cfg = tender_config.get("parser", {})
    client = BidLLMClient(
        model=ai.get("model", "deepseek-v4-flash"),
        base_url=ai.get("base_url", "https://api.deepseek.com"),
    )
    return BidParsePipeline(client, parser_cfg)


def _resolve_mode(project_id: int, mode: Optional[str]) -> str:
    """解析 mode：query 参数 > config 默认。"""
    if mode and mode in {"auto", "quick", "full", "manual"}:
        return mode
    return tender_config.get("parser", {}).get("mode", "auto")


def _get_project_or_404(project_id: int) -> Project:
    session = get_session()
    project = session.get(Project, project_id)
    if not project:
        raise HTTPException(404, "项目不存在")
    tender_path = Path(project.tender_file_path)
    if not tender_path.exists():
        raise HTTPException(400, f"招标文件不存在：{tender_path}")
    return project


def _step1_summary(text: str, info: dict) -> dict:
    return {
        "text_length": len(text),
        "file_type": info.get("type"),
        "pages": info.get("estimated_pages", 0),
    }


@app.post("/projects/{project_id}/parse/step1")
def parse_step1(project_id: int, mode: Optional[str] = None):
    """Step 1 — 文件文本提取（无 LLM，最快 ~1s）。"""
    project = _get_project_or_404(project_id)
    pipeline = _make_pipeline()
    t0 = time.time()
    text = pipeline.step1_extract_text(project.tender_file_path)
    info = get_file_info(project.tender_file_path)
    _parse_state[project_id] = {
        "mode": _resolve_mode(project_id, mode),
        "file_path": project.tender_file_path,
        "file_info": info,
        "full_text": text,
        "t_start": t0,
    }
    return {
        "step": 1, "name": "提取文本", "status": "done",
        "elapsed_sec": round(time.time() - t0, 2),
        "summary": _step1_summary(text, info),
    }


@app.post("/projects/{project_id}/parse/step2")
def parse_step2(project_id: int):
    """Step 2 — 单次 LLM 全量解析（1M 上下文，~20-40s）。

    替代旧的 step2_detect_markers + step3_scan_and_extract + step4_extract_fields 三步。
    1 次 LLM 调用同时输出 K01-K14 + 8 模块 + 标记抽取。
    """
    state = _parse_state.get(project_id)
    if not state:
        raise HTTPException(400, "请先执行 step1")
    if state["mode"] in ("quick", "manual"):
        return {"step": 2, "name": "LLM 解析", "status": "skipped", "summary": {}}
    pipeline = _make_pipeline()
    t0 = time.time()
    parsed = pipeline.step2_full_parse(state["full_text"])
    state["parsed"] = parsed

    if parsed.get("_mode") in ("manual", "error"):
        return {
            "step": 2, "name": "LLM 解析", "status": parsed["_mode"],
            "elapsed_sec": round(time.time() - t0, 2),
            "summary": {"_error": parsed.get("_error", "")},
        }

    k_filled = sum(1 for k in parsed if k.startswith("K") and parsed[k] and parsed[k] != "未找到")
    modules = ["base", "qualification", "rejection", "scoring", "tech", "commercial", "templates", "logistics"]
    modules_filled = [m for m in modules if parsed.get(m)]
    marker_fatal = len(parsed.get("marker_extractions", {}).get("fatal_items", []))
    marker_total = parsed.get("marker_extractions", {}).get("extraction_summary", {}).get("total_marker_occurrences", 0)

    return {
        "step": 2, "name": "LLM 解析", "status": "done",
        "elapsed_sec": round(time.time() - t0, 2),
        "summary": {
            "k_filled": k_filled,
            "modules": modules,
            "modules_filled": modules_filled,
            "marker_fatal": marker_fatal,
            "marker_total": marker_total,
        },
    }


@app.post("/projects/{project_id}/parse/step3")
def parse_step3(project_id: int):
    """Step 3 — 本地校验合并 + 落库。

    接收 step2 的 LLM 输出，做本地校验（模块完整性、必填字段），
    落库到 Project.parsed_data，更新 status='parsed'。
    """
    state = _parse_state.get(project_id)
    if not state:
        raise HTTPException(400, "请先执行 step1")
    t0 = time.time()
    project = _get_project_or_404(project_id)
    mode = state["mode"]

    if mode == "manual":
        # 无 LLM 可用：只跑文本提取 + 标记扫描
        from agents.bid_parser.marker_scanner import (
            extract_pages as _extract_pages,
            scan_markers as _scan_markers,
            summarize_markers as _summarize_markers,
        )
        pages = _extract_pages(state["full_text"])
        hits = _scan_markers(pages)
        marker_sum = _summarize_markers(hits)
        output = {
            "_mode": "manual",
            "meta": {"text_length": len(state["full_text"])},
            "_text_length": len(state["full_text"]),
            "_text_preview": state["full_text"][:2000],
            "_marker_summary": marker_sum,
        }
    elif mode == "quick":
        pipeline = _make_pipeline()
        result = pipeline.run_quick(state["file_path"])
        output = {**result, "_mode": "quick", "_text_length": len(state["full_text"])}
    else:
        # full / auto：合并 step2 输出 + 本地校验
        if "parsed" not in state:
            raise HTTPException(400, "full 模式需要 step2 完成")
        parsed = state["parsed"]
        if parsed.get("_mode") in ("manual", "error"):
            output = {**parsed, "_text_length": len(state["full_text"])}
        else:
            pipeline = _make_pipeline()
            final = pipeline.step3_validate(parsed, state["full_text"])
            # K-层提取（LLM 已给，fallback 用 formatters）
            from agents.parser_agent import ParserAgent
            k01_k14 = ParserAgent({})._extract_k01_k14(final)
            output = {
                **k01_k14,
                **{k: v for k, v in final.items() if k not in k01_k14},
                "_mode": "full",
                "_text_length": len(state["full_text"]),
            }

    # 落库
    session = get_session()
    p = session.get(Project, project_id)
    p.parsed_data = json.dumps(output, ensure_ascii=False)
    p.status = "parsed"
    session.commit()

    # 清理缓存
    _parse_state.pop(project_id, None)

    total_elapsed = round(time.time() - state.get("t_start", t0), 2)
    return {
        "step": 3, "name": "校验合并", "status": "done",
        "elapsed_sec": round(time.time() - t0, 2),
        "total_elapsed_sec": total_elapsed,
        "summary": {
            "validation_issues": len((output.get("_validation") or {}).get("issues", [])),
        },
        "parsed_data": output,
    }


# ============================================================
# 流式解析（AI SDK Data Stream Protocol）— 阶段 2
# 旧 /parse/step1/2/3 端点保留以兼容 Vue 前端，迁移完成后删除
# ============================================================

from sse_stream import (
    stream_text as _sse_text,
    stream_tool_call as _sse_tool,
    stream_finish as _sse_finish,
    stream_error as _sse_error,
)


async def _run_parse_sse(project_id: int, mode: str):
    """
    流式解析生成器：3 步管道 emit AI SDK Protocol 事件。

    实现：worker thread 把事件推到 queue，async generator 从 queue 拉。
    这样 LLM 同步流式调用不会阻塞 FastAPI event loop。
    """
    import asyncio
    import queue
    import threading

    q: queue.Queue = queue.Queue()
    SENTINEL = object()

    def _worker():
        """运行在工作线程：执行 3 步管道，把事件推入 queue。"""
        try:
            pipeline = _make_pipeline()

            # 加载项目
            session = get_session()
            project = session.get(Project, project_id)
            if not project:
                for ev in _sse_error_sync("项目不存在"):
                    q.put(ev)
                return
            tender_path = Path(project.tender_file_path)
            if not tender_path.exists():
                for ev in _sse_error_sync(f"招标文件不存在：{tender_path}"):
                    q.put(ev)
                return

            # ---- Step 1: 提取文本 ----
            t0 = time.time()
            for ev in _sse_text_sync("📄 步骤 1/3 — 提取文本..."):
                q.put(ev)
            text = pipeline.step1_extract_text(tender_path)
            info = get_file_info(tender_path)
            step1_elapsed = round(time.time() - t0, 2)
            for ev in _sse_text_sync(
                f"✓ 提取完成：{len(text)} 字符，{info.get('estimated_pages', 0)} 页（{step1_elapsed}s）"
            ):
                q.put(ev)

            # ---- Step 2: 流式 LLM 解析 ----
            t0 = time.time()
            for ev in _sse_text_sync("🤖 步骤 2/3 — 调用 LLM（最多 16K tokens，~20-40s）..."):
                q.put(ev)

            if mode == "manual" or not pipeline.llm.is_available:
                # 无 LLM 可用 — 走降级路径
                for ev in _sse_text_sync("⚠️ LLM 不可用，进入 manual 模式（仅文件提取 + 标记扫描）"):
                    q.put(ev)
                from agents.bid_parser.marker_scanner import (
                    extract_pages as _extract_pages,
                    scan_markers as _scan_markers,
                    summarize_markers as _summarize_markers,
                )
                pages = _extract_pages(text)
                hits = _scan_markers(pages)
                marker_sum = _summarize_markers(hits)
                output = {
                    "_mode": "manual",
                    "meta": {"text_length": len(text)},
                    "_text_length": len(text),
                    "_text_preview": text[:2000],
                    "_marker_summary": marker_sum,
                }
                # 落库
                p = session.get(Project, project_id)
                p.parsed_data = json.dumps(output, ensure_ascii=False)
                p.status = "parsed"
                session.commit()
                for ev in _sse_tool_sync(
                    "parseTender",
                    {"mode": mode, "projectId": project_id},
                    output,
                ):
                    q.put(ev)
                for ev in _sse_finish_sync():
                    q.put(ev)
                return

            # full / quick 模式 — 调用 LLM（不流式输出原始 JSON，前端只关心最终结果）
            accumulated_text: list[str] = []
            for delta in pipeline.step2_full_parse_stream(text):
                accumulated_text.append(delta)

            step2_elapsed = round(time.time() - t0, 2)
            full_text = "".join(accumulated_text)
            for ev in _sse_text_sync(
                f"✓ LLM 解析完成：{len(full_text)} 字符，{step2_elapsed}s"
            ):
                q.put(ev)

            # 解析 LLM 输出为 dict
            # 鲁棒 JSON 提取：LLM 不带 response_format 后可能输出
            #   - 纯 JSON
            #   - ```json\n{...}\n``` fence
            #   - "Here's the JSON: ... { ... }" 带前缀
            # _extract_json_object 做括号配对扫描，避开字符串内花括号
            parsed = _extract_json_object(full_text.strip())

            if not parsed:
                # 兜底：fence 剥离后重试
                cleaned = full_text.strip()
                if cleaned.startswith("```"):
                    cleaned = re.sub(r"^```(?:json)?\s*\n?", "", cleaned)
                    cleaned = re.sub(r"\n?```\s*$", "", cleaned)
                    parsed = _extract_json_object(cleaned)

            if not parsed:
                # 流式输出 JSON 截断 / 格式异常 → 兜底走非流式（内置 2 次重试）
                for ev in _sse_text_sync("⚠️ 流式输出解析失败，自动重试…"):
                    q.put(ev)
                fallback = pipeline.step2_full_parse(text)
                if fallback and "_error" not in fallback:
                    parsed = fallback
                    for ev in _sse_text_sync("✓ 重试成功"):
                        q.put(ev)
                else:
                    err_msg = (fallback or {}).get("_error", "LLM 输出无法解析为 JSON")
                    for ev in _sse_error_sync(err_msg):
                        q.put(ev)
                    # 把原文前 300 字符回传前端方便诊断
                    preview = full_text[:300].replace("\n", " ")
                    for ev in _sse_text_sync(f"原始输出片段: {preview}…"):
                        q.put(ev)
                    return

            # 注入 meta
            parsed.setdefault("meta", {})
            parsed["meta"]["parser_version"] = "3.1.0-full-context"
            parsed["meta"]["text_length"] = len(text)
            parsed["meta"]["prompt_mode"] = "single-shot-1M-context"

            # ---- Step 3: 校验合并 + 落库 ----
            t0 = time.time()
            for ev in _sse_text_sync("✅ 步骤 3/3 — 校验合并..."):
                q.put(ev)
            final = pipeline.step3_validate(parsed, text)
            from agents.parser_agent import ParserAgent
            k01_k14 = ParserAgent({})._extract_k01_k14(final)
            output = {
                **k01_k14,
                **{k: v for k, v in final.items() if k not in k01_k14},
                "_mode": "full",
                "_text_length": len(text),
            }
            p = session.get(Project, project_id)
            p.parsed_data = json.dumps(output, ensure_ascii=False)
            p.status = "parsed"
            session.commit()
            step3_elapsed = round(time.time() - t0, 2)
            for ev in _sse_text_sync(f"✓ 已落库，{step3_elapsed}s"):
                q.put(ev)

            for ev in _sse_tool_sync(
                "parseTender",
                {"mode": mode, "projectId": project_id},
                output,
            ):
                q.put(ev)
            for ev in _sse_finish_sync():
                q.put(ev)

        except Exception as e:
            for ev in _sse_error_sync(f"解析异常：{type(e).__name__}: {e}"):
                q.put(ev)
        finally:
            q.put(SENTINEL)

    # 启动 worker 线程
    t = threading.Thread(target=_worker, daemon=True)
    t.start()

    # async generator 从 queue 拉事件
    loop = asyncio.get_event_loop()
    while True:
        ev = await loop.run_in_executor(None, q.get)
        if ev is SENTINEL:
            break
        yield ev


# 同步版本的 SSE 事件生成器（用于 worker 线程）
def _sse_text_sync(text, message_id=None, text_id=None):
    import uuid as _uuid
    message_id = message_id or f"msg_{_uuid.uuid4().hex[:12]}"
    text_id = text_id or f"text_{_uuid.uuid4().hex[:12]}"
    yield {"data": json.dumps({"type": "start", "messageId": message_id}, ensure_ascii=False)}
    yield {"data": json.dumps({"type": "text-start", "id": text_id}, ensure_ascii=False)}
    yield {"data": json.dumps({"type": "text-delta", "id": text_id, "delta": text}, ensure_ascii=False)}
    yield {"data": json.dumps({"type": "text-end", "id": text_id}, ensure_ascii=False)}
    yield {"data": json.dumps({"type": "finish-step"}, ensure_ascii=False)}


def _sse_tool_sync(tool_name, tool_input, tool_output, message_id=None, tool_call_id=None):
    import uuid as _uuid
    message_id = message_id or f"msg_{_uuid.uuid4().hex[:12]}"
    tool_call_id = tool_call_id or f"call_{_uuid.uuid4().hex[:12]}"
    yield {"data": json.dumps({"type": "start", "messageId": message_id}, ensure_ascii=False)}
    yield {"data": json.dumps(
        {"type": "tool-input-available", "toolCallId": tool_call_id, "toolName": tool_name, "input": tool_input},
        ensure_ascii=False,
    )}
    yield {"data": json.dumps(
        {"type": "tool-output-available", "toolCallId": tool_call_id, "output": tool_output},
        ensure_ascii=False,
    )}
    yield {"data": json.dumps({"type": "finish-step"}, ensure_ascii=False)}


def _extract_json_object(text: str) -> Optional[dict]:
    """
    从 LLM 输出里抽第一个完整的 JSON 对象。

    鲁棒处理：前缀文字、markdown fence、字符串内花括号、嵌套结构。
    顶屋必须是 dict（招标解析 schema 是 object）。
    """
    if not text:
        return None
    fence = re.search(r"```(?:json)?\s*\n(.*?)\n?```", text, re.DOTALL)
    if fence:
        text = fence.group(1).strip()
    start_idx = -1
    opener, closer = "", ""
    for i, ch in enumerate(text):
        if ch == "{":
            start_idx, opener, closer = i, "{", "}"
            break
        if ch == "[":
            start_idx, opener, closer = i, "[", "]"
            break
    if start_idx < 0:
        return None
    depth = 0
    in_string = escape = False
    for j in range(start_idx, len(text)):
        c = text[j]
        if escape:
            escape = False
            continue
        if c == "\\":
            escape = True
            continue
        if c == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if c == opener:
            depth += 1
        elif c == closer:
            depth -= 1
            if depth == 0:
                try:
                    obj = json.loads(text[start_idx:j + 1])
                    if isinstance(obj, dict):
                        return obj
                except json.JSONDecodeError:
                    return None
                return None
    return None


def _sse_finish_sync():
    yield {"data": json.dumps({"type": "finish"}, ensure_ascii=False)}


def _sse_error_sync(error_message):
    yield {"data": json.dumps({"type": "error", "errorText": error_message}, ensure_ascii=False)}


def _sse_readiness_failure_sync(tool_name: str, readiness: dict[str, Any], tool_input: dict[str, Any]):
    message = (
        f"{readiness.get('message')} 缺少："
        f"{'、'.join(readiness.get('missing') or [])}。"
        f"{readiness.get('recoverable_action')}"
    )
    yield from _sse_tool_sync(
        tool_name,
        tool_input,
        {
            "error": "readiness_failed",
            "message": message,
            "missing": readiness.get("missing", []),
            "requires": readiness.get("requires", []),
            "recoverable_action": readiness.get("recoverable_action"),
        },
    )
    yield from _sse_finish_sync()


@app.get("/projects/{project_id}/parse/stream")
async def parse_stream(project_id: int, mode: Optional[str] = None):
    """
    流式解析端点 — emit AI SDK Data Stream Protocol。

    替换旧的 /parse/step1+2+3 串行调用。前端 useChat 直连。

    Query:
        mode: auto/quick/full/manual（缺省走 config.yaml）
    """
    resolved = _resolve_mode(project_id, mode)
    return EventSourceResponse(_run_parse_sse(project_id, resolved))


# ============================================================
# Chat 端点 — useChat 的标准入口，关键词路由
# ============================================================


async def _run_chat_sse(project_id: int, last_user_msg: str):
    """
    关键词路由：把 useChat 的消息分派到对应 agent 流程。

    路由规则（按出现顺序匹配首个）：
      "放好了" / "解析"  → parse 流（delegate to _run_parse_sse）
      "继续"             → match
      "生成" / "开始生成" → generate
      "终审"             → review
      "导出" / "下载"     → export
      else               → help / echo
    """
    msg = last_user_msg.strip()

    session = get_session()
    project = session.get(Project, project_id)
    if not project:
        for ev in _sse_error_sync("项目不存在"):
            yield ev
        return

    # ---- "放好了" → parse ----
    if _should_start_parse_from_message(msg, project.status):
        readiness = _check_tool_readiness(project, "parseTender", session)
        if not readiness["ok"]:
            for ev in _sse_readiness_failure_sync(
                "parseTender",
                readiness,
                {"projectId": project_id, "mode": "full"},
            ):
                yield ev
            return
        for ev in _sse_text_sync("🔍 收到，开始解析..."):
            yield ev
        async for ev in _run_parse_sse(project_id, "full"):
            yield ev
        return

    # ---- parsed-data natural-language corrections ----
    correction_patches = (
        _extract_parsed_data_patch_proposals(
            msg,
            json.loads(project.parsed_data) if project.parsed_data else {},
        )
        if project.parsed_data else []
    )
    if correction_patches:
        try:
            correction_result = _apply_parsed_data_patches(project, correction_patches)
        except ValueError as e:
            for ev in _sse_error_sync(f"解析修正失败：{e}"):
                yield ev
            return
        session.commit()
        for ev in _sse_text_sync(correction_result["message"]):
            yield ev
        for ev in _sse_tool_sync(
            "parseTender",
            {"projectId": project_id, "source": "chat_nl_correction"},
            {
                **correction_result["parsed_data"],
                "_correction_applied": correction_result["applied"],
                "message": correction_result["message"],
                "correction_hint": "修正已写入解析数据；确认无误后说「继续」生成提纲。",
            },
        ):
            yield ev
        for ev in _sse_finish_sync():
            yield ev
        return

    # ---- price calculation tool ----
    if _should_run_price_calculator(msg):
        parsed = json.loads(project.parsed_data) if project.parsed_data else {}
        inputs = _extract_price_calc_inputs(msg, parsed)
        result = _calculate_price_scores(parsed, **inputs)
        for ev in _sse_tool_sync(
            "priceCalculator",
            {"projectId": project_id, **inputs},
            {
                **result,
                "inputs": inputs,
                "action_hint": (
                    "可补充“低价比例80%、最高限价900万、价格分30分、最低报价720万、主标报价760万、对手报价780万”。"
                    if not result.get("ok")
                    else "如要复算，可直接给新的最低报价、主标报价和竞争对手报价。"
                ),
            },
        ):
            yield ev
        for ev in _sse_finish_sync():
            yield ev
        return

    # ---- "继续" / 其他 → 走 Orchestrator 状态机（关键路径）----
    # 状态机驱动：AWAIT_PARSE_CONFIRM → generate_outline → AWAIT_OUTLINE_CONFIRM
    #              AWAIT_OUTLINE_CONFIRM → match_materials → AWAIT_CHAPTER_CONFIRM
    #              AWAIT_CHAPTER_CONFIRM → generate → AWAIT_DRAFT_CONFIRM
    # handler 返回 dict 里有 outline / chapters / draft_preview 键就发对应 tool 事件
    # 上面分支已 return（parse / generate / review / export 走一次性路径），下面就是状态机入口
    # ---- "终审" → review ----
    if "终审" in msg:
        readiness = _check_tool_readiness(project, "reviewTender", session)
        if not readiness["ok"]:
            for ev in _sse_readiness_failure_sync(
                "reviewTender",
                readiness,
                {"projectId": project_id, "tenderType": "main"},
            ):
                yield ev
            return
        tender = _active_main_tender(session, project_id)
        if not tender:
            for ev in _sse_error_sync("未找到主标书，无法终审"):
                yield ev
            return
        for ev in _sse_text_sync("🔍 终审检查中..."):
            yield ev
        orch = Orchestrator(tender_config)
        orch.ctx.tender_id = tender.id
        orch.ctx.tender_type = "main"
        orch.ctx.project_id = project_id
        review_result = orch.agents["reviewer"].execute(orch.ctx)
        next_status = "review_failed" if review_result.get("error") else "reviewed"
        project.status = next_status
        tender.status = next_status
        session.commit()
        for ev in _sse_tool_sync(
            "reviewTender",
            {"tenderId": tender.id},
            {
                "tenderId": tender.id,
                "checks": review_result.get("checks", []),
                "issues": review_result.get("issues", []),
                "summary": review_result.get("summary", {}),
                "deterministic_count": review_result.get("deterministic_count", 0),
                "error": review_result.get("error"),
                "message": (
                    f"终审失败：{review_result.get('error')}"
                    if review_result.get("error")
                    else "终审检查完成"
                ),
                "action_hint": (
                    "请先修复终审前置条件后重试"
                    if review_result.get("error")
                    else "要我一键修正，还是你手动处理？"
                ),
            },
        ):
            yield ev
        for ev in _sse_finish_sync():
            yield ev
        return

    # ---- "导出提纲" → outline export ----
    if ("导出" in msg or "下载" in msg) and re.search(r"提纲|大纲|目录", msg):
        readiness = _check_tool_readiness(project, "exportOutline", session)
        if not readiness["ok"]:
            for ev in _sse_readiness_failure_sync(
                "exportTender",
                readiness,
                {"projectId": project_id, "format": "markdown", "target": "outline"},
            ):
                yield ev
            return
        try:
            export_result = _export_outline_file(project, "markdown")
        except HTTPException as e:
            for ev in _sse_tool_sync(
                "exportTender",
                {"projectId": project_id, "format": "markdown", "target": "outline"},
                {
                    "projectId": project_id,
                    "format": "markdown",
                    "error": e.detail,
                    "message": f"提纲导出失败：{e.detail}",
                },
            ):
                yield ev
            for ev in _sse_finish_sync():
                yield ev
            return
        for ev in _sse_tool_sync(
            "exportTender",
            {"projectId": project_id, "format": "markdown", "target": "outline"},
            {
                "projectId": project_id,
                **export_result,
                "message": "已导出提纲为 markdown",
            },
        ):
            yield ev
        for ev in _sse_finish_sync():
            yield ev
        return

    # ---- "导出" → export ----
    if "导出" in msg or "下载" in msg:
        readiness = _check_tool_readiness(project, "exportTender", session)
        if not readiness["ok"]:
            for ev in _sse_readiness_failure_sync(
                "exportTender",
                readiness,
                {"projectId": project_id, "format": _export_format_from_message(msg)},
            ):
                yield ev
            return
        tender = _active_main_tender(session, project_id)
        if not tender:
            for ev in _sse_error_sync("未找到主标书，无法导出"):
                yield ev
                return
        export_format = _export_format_from_message(msg)
        try:
            export_result = _export_tender_file(tender, export_format)
        except HTTPException as e:
            for ev in _sse_tool_sync(
                "exportTender",
                {"tenderId": tender.id, "format": export_format},
                {
                    "tenderId": tender.id,
                    "format": export_format,
                    "error": e.detail,
                    "message": f"导出失败：{e.detail}",
                },
            ):
                yield ev
            for ev in _sse_finish_sync():
                yield ev
            return
        for ev in _sse_tool_sync(
            "exportTender",
            {"tenderId": tender.id, "format": export_result["format"]},
            {
                "tenderId": tender.id,
                **export_result,
                "message": f"已导出为 {export_result['format']}",
            },
        ):
            yield ev
        for ev in _sse_finish_sync():
            yield ev
        return

    # 用 project.status 恢复 orch.step（避免依赖全局 .session.json）
    status_to_step = {
        "parsed": WorkflowStep.AWAIT_PARSE_CONFIRM,
        "outline_generating": WorkflowStep.AWAIT_OUTLINE_CONFIRM,
        "materials_preparing": WorkflowStep.AWAIT_CHAPTER_CONFIRM,
        "deviation_preparing": WorkflowStep.AWAIT_DEVIATION_CONFIRM,
        "generating": WorkflowStep.AWAIT_DRAFT_CONFIRM,
        "generated": WorkflowStep.AWAIT_DRAFT_CONFIRM,
        "reviewing": WorkflowStep.AWAIT_REVIEW_ACTION,
        "reviewed": WorkflowStep.AWAIT_REVIEW_ACTION,
        "review_failed": WorkflowStep.AWAIT_REVIEW_ACTION,
        "done": WorkflowStep.AWAIT_REVIEW_ACTION,
    }

    orch = Orchestrator(tender_config)
    # 恢复 ctx
    orch.ctx.project_id = project_id
    orch.ctx.parsed_data = json.loads(project.parsed_data) if project.parsed_data else {}
    # 恢复 outline / chapters from parsed_data
    outline = (orch.ctx.parsed_data.get("_confirmed_outline")
               or orch.ctx.parsed_data.get("_generated_outline") or [])
    if outline:
        orch.ctx.outline = outline
    if "chapters" in orch.ctx.parsed_data:
        orch.ctx.chapters = orch.ctx.parsed_data["chapters"]
    # 强制设置 step 为当前 project 状态
    orch.step = status_to_step.get(project.status, WorkflowStep.IDLE)

    state_tool = None
    if "继续" in msg:
        state_tool = {
            "parsed": "outlineDesign",
            "outline_generating": "matchMaterials",
            "materials_preparing": "generateTender",
            "deviation_preparing": "generateTender",
        }.get(project.status)
    elif re.search(r"生成|开始生成", msg):
        state_tool = "generateTender"
    if state_tool:
        readiness = _check_tool_readiness(project, state_tool, session)
        if not readiness["ok"]:
            for ev in _sse_readiness_failure_sync(
                state_tool,
                readiness,
                {"projectId": project_id, "tenderType": "main"},
            ):
                yield ev
            return

    # "继续" 状态机的"进入动画"文案
    if "继续" in msg or "生成" in msg:
        if project.status == "parsed":
            for ev in _sse_text_sync("📑 生成章节大纲中..."):
                yield ev
        elif project.status == "outline_generating":
            for ev in _sse_text_sync("🔍 匹配材料中..."):
                yield ev
        elif project.status == "materials_preparing":
            for ev in _sse_text_sync("📋 准备偏离表确认项..."):
                yield ev
        elif project.status == "deviation_preparing":
            for ev in _sse_text_sync("📝 生成标书中（多 Agent 串行，可能需要数分钟）..."):
                yield ev
        elif "陪标" in msg:
            for ev in _sse_text_sync("📋 生成陪标并执行终审中..."):
                yield ev

    # 让 Orchestrator 状态机处理
    result = orch.handle(msg)
    message = result.get("message", "")

    # 按返回 dict 的 key 发对应 tool 事件。
    # 关键：发完 tool 事件后不要再把同一段 message 文本再发一次 text-delta —
    # 前端的 tool 组件（如 OutlineToolResult）已经渲染了同样的内容，重复发就是 UI 上
    # 看到两次。
    emitted_tool = False
    if "outline" in result and "chapters" not in result:
        for ev in _sse_tool_sync(
            "outlineDesign",
            {"projectId": project_id, "tenderType": "main"},
            {
                "outline": result["outline"],
                "message": message,
                "action_hint": "说「继续」进入材料匹配；也可以直接告诉我怎么改，比如「把第3章删了」「加一章 数据迁移」「重新生成」",
            },
        ):
            yield ev
        emitted_tool = True
    if "chapters" in result and "outline" not in result:
        for ev in _sse_tool_sync(
            "matchMaterials",
            {"projectId": project_id, "tenderType": "main"},
            {
                "chapters": result["chapters"],
                "message": message,
                "action_hint": "可以告诉我需要替换哪个章节，或说'继续'进入生成。",
            },
        ):
            yield ev
        emitted_tool = True
    if "deviation_preview" in result:
        for ev in _sse_tool_sync(
            "confirmDeviation",
            {"projectId": project_id, "tenderType": "main"},
            {
                **result["deviation_preview"],
                "message": message,
                "action_hint": "确认无误请说「确认」或「继续」；需要补充时说「补充商务偏离：...」或「补充技术偏离：...」。",
            },
        ):
            yield ev
        emitted_tool = True
    if "draft_preview" in result:
        for ev in _sse_tool_sync(
            "generateTender",
            {"projectId": project_id, "tenderType": "main"},
            {
                "tenderId": result.get("tender_id"),
                "draft_preview": result["draft_preview"],
                "draft_path": result.get("draft_path"),
                "draft_chapters": result.get("draft_chapters", []),
                "errors": result.get("errors", []),
                "failed": bool(result.get("failed")),
                "outline": result.get("outline", []),
                "message": message,
            },
        ):
            yield ev
        emitted_tool = True
    if "sub_draft" in result:
        sub_review = result.get("sub_review") or {}
        for ev in _sse_tool_sync(
            "generateTender",
            {"projectId": project_id, "tenderType": "sub"},
            {
                "tenderId": result.get("tender_id"),
                "draft_preview": (result.get("sub_draft") or {}).get("content", "")[:500],
                "draft_path": result.get("draft_path"),
                "draft_chapters": [],
                "errors": (result.get("sub_draft") or {}).get("errors", []),
                "failed": bool(result.get("failed")),
                "message": message,
            },
        ):
            yield ev
        for ev in _sse_tool_sync(
            "reviewTender",
            {"tenderId": result.get("tender_id"), "tenderType": "sub"},
            {
                "tenderId": result.get("tender_id"),
                "tenderType": "sub",
                "checks": sub_review.get("checks", []),
                "issues": sub_review.get("issues", []),
                "summary": sub_review.get("summary", {}),
                "deterministic_count": sub_review.get("deterministic_count", 0),
                "error": sub_review.get("error"),
                "retries": result.get("retries", 0),
                "message": "陪标终审完成",
            },
        ):
            yield ev
        emitted_tool = True

    # 没有 tool 事件 → 当成纯文本消息（如 help、unknown command）
    if not emitted_tool and message:
        for ev in _sse_text_sync(message):
            yield ev
    for ev in _sse_finish_sync():
        yield ev
    return


@app.post("/projects/{project_id}/chat")
async def chat_router(project_id: int, request: Request):
    """
    useChat 主入口 — 接收 messages[]，按关键词路由到对应 agent 流。

    接收：
        { "messages": [{role, content}, ...] }

    返回：
        SSE 流（AI SDK Data Stream Protocol）

    替代旧 /projects/{id}/parse/confirm + /generate + /tenders/{id}/review + /tenders/{id}/export 的
    多端点跳转模式——所有动作统一通过这一个 chat 端点触发。
    """
    try:
        body = await request.json()
    except Exception:
        raise HTTPException(400, "请求体必须为 JSON")
    messages = body.get("messages", [])
    if not isinstance(messages, list):
        raise HTTPException(400, "messages 必须为数组")
    # 找最后一条 user 消息
    # AI SDK v3: { role: "user", parts: [{ type: "text", text: "..." }] }
    # v2 兼容: { role: "user", content: "..." } 或 content: [{ type: "text", text: "..." }]
    last_user_msg = ""
    for m in reversed(messages):
        if isinstance(m, dict) and m.get("role") == "user":
            # v3 优先
            parts = m.get("parts")
            if isinstance(parts, list) and parts:
                last_user_msg = "".join(
                    p.get("text", "") for p in parts if isinstance(p, dict) and p.get("type") == "text"
                )
            else:
                # v2 fallback
                content = m.get("content", "")
                if isinstance(content, str):
                    last_user_msg = content
                elif isinstance(content, list):
                    last_user_msg = "".join(
                        p.get("text", "") for p in content if isinstance(p, dict) and p.get("type") == "text"
                    )
            break

    # 校验项目存在
    session = get_session()
    project = session.get(Project, project_id)
    if not project:
        raise HTTPException(404, "项目不存在")

    return EventSourceResponse(_run_chat_sse(project_id, last_user_msg))


@app.post("/projects/{project_id}/parse/confirm")
def confirm_parse(project_id: int, req: TenderParseConfirmRequest):
    """用户确认解析结果，一次性生成提纲 + 匹配材料（Web 走一次性路径）。"""
    session = get_session()
    project = session.get(Project, project_id)
    if not project:
        raise HTTPException(404, "项目不存在")

    if req.corrections:
        for key, value in req.corrections.items():
            if hasattr(project, key):
                setattr(project, key, value)

    project.status = "materials_preparing"
    session.commit()

    orch = Orchestrator(tender_config)
    orch.ctx.project_id = project_id
    orch.ctx.parsed_data = json.loads(project.parsed_data) if project.parsed_data else {}

    # 一次性调用（Web 跳过中间确认步骤）
    match_result = orch.agents["matcher"].execute(orch.ctx)
    orch.ctx.parsed_data["_confirmed_outline"] = match_result.get("outline", [])
    orch.ctx.parsed_data["chapters"] = match_result.get("chapters", [])

    project.parsed_data = json.dumps(orch.ctx.parsed_data, ensure_ascii=False)
    session.commit()

    return {
        "message": "信息已确认，材料匹配完成",
        "outline": match_result.get("outline", []),
        "chapters": match_result.get("chapters", []),
        "next_action": "开始主标材料匹配"
    }


def _parse_field_path(path: str) -> list[str | int]:
    parts: list[str | int] = []
    for seg in path.split("."):
        if not seg:
            raise ValueError("字段路径不能为空")
        pos = 0
        name = ""
        while pos < len(seg):
            ch = seg[pos]
            if ch == "[":
                if name:
                    parts.append(name)
                    name = ""
                end = seg.find("]", pos)
                if end < 0:
                    raise ValueError(f"字段路径格式错误: {path}")
                idx_text = seg[pos + 1:end]
                if not idx_text.isdigit():
                    raise ValueError(f"数组下标必须是数字: {path}")
                parts.append(int(idx_text))
                pos = end + 1
            else:
                name += ch
                pos += 1
        if name:
            parts.append(name)
    return parts


def _get_path_value(data: Any, parts: list[str | int]) -> Any:
    cur = data
    for part in parts:
        if isinstance(part, int):
            if not isinstance(cur, list) or part >= len(cur):
                return None
            cur = cur[part]
        else:
            if not isinstance(cur, dict) or part not in cur:
                return None
            cur = cur[part]
    return cur


def _set_path_value(data: dict, field_path: str, value: Any) -> Any:
    """Set a parsed_data path and return the old value."""
    if not field_path or field_path.startswith("_"):
        raise ValueError("不允许修改内部字段")

    # K 字段顶层修改保留现有 source_page/source_pages。
    if "." not in field_path and "[" not in field_path and re.match(r"^K\d{2}_", field_path):
        old = data.get(field_path)
        if isinstance(old, dict) and ("value" in old or "items" in old):
            if isinstance(value, list):
                pages = old.get("source_pages") if isinstance(old.get("source_pages"), list) else []
                data[field_path] = {"items": value, "source_pages": pages}
            else:
                data[field_path] = {
                    "value": value,
                    "source_page": old.get("source_page"),
                }
        else:
            data[field_path] = value
        return old

    parts = _parse_field_path(field_path)
    old = _get_path_value(data, parts)
    cur: Any = data
    for i, part in enumerate(parts[:-1]):
        nxt = parts[i + 1]
        if isinstance(part, int):
            if not isinstance(cur, list):
                raise ValueError(f"路径不是数组: {field_path}")
            while len(cur) <= part:
                cur.append({} if isinstance(nxt, str) else [])
            cur = cur[part]
        else:
            if not isinstance(cur, dict):
                raise ValueError(f"路径不是对象: {field_path}")
            if part not in cur or cur[part] is None:
                cur[part] = [] if isinstance(nxt, int) else {}
            cur = cur[part]

    last = parts[-1]
    if isinstance(last, int):
        if not isinstance(cur, list):
            raise ValueError(f"路径不是数组: {field_path}")
        while len(cur) <= last:
            cur.append(None)
        cur[last] = value
    else:
        if not isinstance(cur, dict):
            raise ValueError(f"路径不是对象: {field_path}")
        cur[last] = value
    return old


def _source_page_patch_for_field(
    parsed: dict[str, Any],
    field_path: str,
    page: int,
    note: str,
) -> ParsedDataPatchItem:
    field = parsed.get(field_path) if isinstance(parsed, dict) else None
    if (
        (isinstance(field, dict) and "items" in field)
        or field_path in ("K10_星标项", "K11_废标条款")
    ):
        items = field.get("items") if isinstance(field, dict) else []
        count = len(items) if isinstance(items, list) and items else 1
        return ParsedDataPatchItem(
            field_path=f"{field_path}.source_pages",
            value=[page] * count,
            note=note,
            source="chat_nl_page",
        )
    return ParsedDataPatchItem(
        field_path=f"{field_path}.source_page",
        value=page,
        note=note,
        source="chat_nl_page",
    )


def _extract_source_page_patch_proposals(
    msg: str,
    parsed: dict[str, Any] | None = None,
) -> list[ParsedDataPatchItem]:
    text = (msg or "").strip()
    if not text:
        return []
    parsed = parsed or {}
    patches: list[ParsedDataPatchItem] = []
    seen: set[str] = set()
    page_expr = r"(?:来源页|原文页|页码|所在页|在)\s*(?:是|为|改为|修改为)?\s*(?:第)?\s*(\d{1,4})\s*页?"

    def add(field_path: str, page_text: str):
        if field_path in seen:
            return
        page = int(page_text)
        if page <= 0:
            return
        seen.add(field_path)
        patches.append(_source_page_patch_for_field(
            parsed,
            field_path,
            page,
            f"自然语言页码修正：{text}",
        ))

    for m in re.finditer(rf"\b(K\d{{2}}_[\w\u4e00-\u9fff]+)\b[^，。,；;\n]{{0,20}}{page_expr}", text):
        add(m.group(1), m.group(2))

    for alias, field_path in sorted(
        CORRECTION_FIELD_ALIASES.items(),
        key=lambda item: len(item[0]),
        reverse=True,
    ):
        m = re.search(rf"{re.escape(alias)}[^，。,；;\n]{{0,20}}{page_expr}", text)
        if m:
            add(field_path, m.group(1))

    return patches


def _extract_parsed_data_patch_proposals(
    msg: str,
    parsed: dict[str, Any] | None = None,
) -> list[ParsedDataPatchItem]:
    """
    Convert common natural-language parsed-data corrections into explicit patches.

    This is deliberately deterministic: it handles the high-frequency K-field
    corrections users make after parsing, while the explicit PATCH endpoint
    remains available for arbitrary nested paths.
    """
    text = (msg or "").strip()
    if not text:
        return []

    verb = r"(?:应为|应该是|应改为|改为|修改为|更正为|设置为|设为|是|为|=|：|:)"
    value = r"\s*([^，。,；;\n]+)"
    patches: list[ParsedDataPatchItem] = []
    seen: set[str] = set()
    patches.extend(_extract_source_page_patch_proposals(text, parsed))
    seen.update(p.field_path for p in patches)

    # Explicit K-field path: "K04_预算金额 改为 900万"
    for m in re.finditer(rf"\b(K\d{{2}}_[\w\u4e00-\u9fff]+)\b\s*{verb}{value}", text):
        field_path = m.group(1)
        new_value = m.group(2).strip()
        if field_path not in seen and new_value:
            seen.add(field_path)
            patches.append(ParsedDataPatchItem(
                field_path=field_path,
                value=new_value,
                note=f"自然语言修正：{text}",
                source="chat_nl",
            ))

    for alias, field_path in CORRECTION_FIELD_ALIASES.items():
        if field_path in seen:
            continue
        m = re.search(rf"{re.escape(alias)}\s*{verb}{value}", text)
        if not m:
            continue
        new_value = m.group(1).strip()
        if not new_value:
            continue
        seen.add(field_path)
        patches.append(ParsedDataPatchItem(
            field_path=field_path,
            value=new_value,
            note=f"自然语言修正：{text}",
            source="chat_nl",
        ))

    return patches


def _sync_project_fields_from_patch(project: Project, patch: ParsedDataPatchItem):
    value = str(patch.value).strip() if patch.value is not None else ""
    if patch.field_path == "K01_项目名称":
        project.project_name = value or None
    elif patch.field_path == "K02_招标编号":
        project.tender_no = value or None
    elif patch.field_path == "K04_预算金额":
        amount = re.sub(r"[^\d.]", "", value)
        project.budget = float(amount) if amount else None


def _apply_parsed_data_patches(
    project: Project,
    patches: list[ParsedDataPatchItem],
) -> dict[str, Any]:
    if not patches:
        raise ValueError("patches 不能为空")

    parsed = json.loads(project.parsed_data) if project.parsed_data else {}
    corrections = parsed.setdefault("_corrections", [])
    applied = []

    for patch in patches:
        old_value = _set_path_value(parsed, patch.field_path, patch.value)
        entry = {
            "field_path": patch.field_path,
            "old_value": old_value,
            "new_value": patch.value,
            "note": patch.note,
            "source": patch.source or "user",
            "updated_at": datetime.utcnow().isoformat(),
        }
        corrections.append(entry)
        applied.append(entry)
        _sync_project_fields_from_patch(project, patch)

    project.parsed_data = json.dumps(parsed, ensure_ascii=False)
    project.status = project.status or "parsed"
    return {
        "message": f"已应用 {len(applied)} 项解析修正",
        "applied": applied,
        "parsed_data": parsed,
    }


@app.patch("/projects/{project_id}/parsed-data")
def patch_parsed_data(project_id: int, req: ParsedDataPatchRequest):
    """Apply explicit parsed_data patches and record correction audit entries."""
    if not req.patches:
        raise HTTPException(400, "patches 不能为空")

    session = get_session()
    project = session.get(Project, project_id)
    if not project:
        raise HTTPException(404, "项目不存在")

    try:
        result = _apply_parsed_data_patches(project, req.patches)
    except ValueError as e:
        raise HTTPException(400, str(e))
    session.commit()

    return result


@app.post("/projects/{project_id}/price-calculator")
def calculate_price(project_id: int, req: PriceCalculationRequest):
    """Calculate bid price scores and abnormal-low-price risk."""
    session = get_session()
    project = session.get(Project, project_id)
    if not project:
        raise HTTPException(404, "项目不存在")
    parsed = json.loads(project.parsed_data) if project.parsed_data else {}
    return _calculate_price_scores(
        parsed,
        lowest_price=req.lowest_price,
        main_price=req.main_price,
        competitor_price=req.competitor_price,
        low_price_ratio=req.low_price_ratio,
        highest_limit=req.highest_limit,
        price_score_max=req.price_score_max,
    )


@app.get("/projects/{project_id}/price-calculator/defaults")
def get_price_calculator_defaults(project_id: int):
    """Return price calculator parameters extracted from parsed tender data."""
    session = get_session()
    project = session.get(Project, project_id)
    if not project:
        raise HTTPException(404, "项目不存在")
    parsed = json.loads(project.parsed_data) if project.parsed_data else {}
    return _price_calculator_defaults(parsed)


def _safe_material_filename(filename: str, fallback: str = "material.md") -> str:
    raw_name = Path(filename or fallback).name
    safe = re.sub(r"[^\w一-鿿\-.]", "_", raw_name).strip("._")
    return safe or fallback


def _next_available_path(path: Path) -> Path:
    if not path.exists():
        return path
    stem = path.stem
    suffix = path.suffix
    for idx in range(2, 1000):
        candidate = path.with_name(f"{stem}_{idx}{suffix}")
        if not candidate.exists():
            return candidate
    raise HTTPException(409, f"同名材料过多：{path.name}")


def _extract_uploaded_material_text(filename: str, content: bytes) -> tuple[str, str]:
    suffix = Path(filename or "").suffix.lower()
    if suffix in {".md", ".txt"}:
        return content.decode("utf-8", errors="ignore"), ".md"
    if suffix == ".docx":
        try:
            from docx import Document
            doc = Document(BytesIO(content))
            text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return text, ".docx"
        except Exception:
            return "", ".docx"
    if suffix == ".pdf":
        try:
            import fitz
            doc = fitz.open(stream=content, filetype="pdf")
            text = "\n\n".join(page.get_text("text") for page in doc)
            return text.strip(), ".md"
        except Exception as exc:
            raise HTTPException(400, f"PDF 材料解析失败：{exc}") from exc
    raise HTTPException(400, f"不支持的材料类型：{suffix or '无后缀'}")


def _write_material_file(category: str, filename: str, content: bytes, text: str, storage_suffix: str) -> Path:
    if category not in STANDARD_MATERIAL_CATEGORIES:
        raise HTTPException(400, "材料分类无效")
    category_dir = MATERIALS_DIR / category
    category_dir.mkdir(parents=True, exist_ok=True)
    safe_name = _safe_material_filename(filename)
    if storage_suffix == ".md":
        target = _next_available_path(category_dir / f"{Path(safe_name).stem}.md")
        if not text.strip():
            text = content.decode("utf-8", errors="ignore")
        target.write_text(text.strip() + "\n", encoding="utf-8")
        return target
    target = _next_available_path(category_dir / safe_name)
    target.write_bytes(content)
    return target


def _material_path_from_request(file_path: str) -> Path:
    root = MATERIALS_DIR.resolve()
    target = Path(file_path).expanduser().resolve()
    try:
        target.relative_to(root)
    except ValueError:
        raise HTTPException(403, "材料路径越界")
    if not target.exists() or not target.is_file():
        raise HTTPException(404, "材料不存在")
    return target


def _read_material_text(path: Path) -> str:
    suffix = path.suffix.lower()
    content = path.read_bytes()
    if suffix in {".md", ".txt"}:
        return content.decode("utf-8", errors="ignore")
    if suffix == ".docx":
        try:
            from docx import Document
            doc = Document(BytesIO(content))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as exc:
            raise HTTPException(400, f"DOCX 材料读取失败：{exc}") from exc
    if suffix == ".pdf":
        try:
            import fitz
            doc = fitz.open(stream=content, filetype="pdf")
            return "\n\n".join(page.get_text("text") for page in doc).strip()
        except Exception as exc:
            raise HTTPException(400, f"PDF 材料读取失败：{exc}") from exc
    raise HTTPException(400, "仅支持读取 .md / .txt / .docx / .pdf 材料")


# ---- 材料库管理 ----

@app.get("/materials")
def list_materials(category: Optional[str] = None, keyword: Optional[str] = None):
    """从文件系统加载材料库（不走数据库）"""
    from agents.matcher_agent import MatcherAgent

    matcher = MatcherAgent()
    all_materials = matcher._load_materials_from_disk()

    # 过滤
    filtered = all_materials
    if category:
        filtered = [m for m in filtered if m["category"] == category]
    if keyword:
        filtered = [m for m in filtered if keyword.lower() in m["title"].lower()]

    # 转换为前端期望的格式 + 读取文件大小
    result = []
    for m in filtered:
        file_path = Path(m["file_path"])
        char_count = 0
        if file_path.exists():
            try:
                if file_path.suffix.lower() == ".md":
                    text = file_path.read_text(encoding="utf-8")
                    char_count = len(text)
            except:
                pass

        result.append({
            "id": hash(m["file_path"]) % 100000,  # 伪 ID（前端需要）
            "title": m["title"],
            "category": m["category"],
            "tags": m.get("tags", ""),
            "description": f"来自 {file_path.name}",
            "char_count": char_count,
            "ai_summary": "",
            "version": 1,
        })

    return result


@app.get("/materials/options")
def list_material_options(category: Optional[str] = None, keyword: Optional[str] = None):
    """返回可用于替换匹配结果的材料，包含 file_path。"""
    from agents.matcher_agent import MatcherAgent

    matcher = MatcherAgent()
    all_materials = matcher._load_materials_from_disk()
    filtered = all_materials
    if category:
        filtered = [m for m in filtered if m["category"] == category]
    if keyword:
        filtered = [m for m in filtered if keyword.lower() in m["title"].lower()]

    result = []
    for m in filtered:
        file_path = Path(m["file_path"])
        char_count = 0
        if file_path.exists() and file_path.suffix.lower() in {".md", ".txt"}:
            try:
                char_count = len(file_path.read_text(encoding="utf-8"))
            except Exception:
                char_count = 0
        result.append({
            "title": m["title"],
            "category": m["category"],
            "file_path": m["file_path"],
            "char_count": char_count,
        })
    return result


@app.get("/materials/read")
def read_material_file(file_path: str):
    path = _material_path_from_request(file_path)
    return {
        "title": path.stem,
        "file_path": str(path),
        "content": _read_material_text(path),
    }


@app.post("/materials")
def create_material(req: CreateMaterialRequest):
    text = req.content.strip()
    if req.category not in STANDARD_MATERIAL_CATEGORIES:
        raise HTTPException(400, "材料分类无效")
    disk_path = _write_material_file(
        req.category,
        f"{req.title or 'material'}.md",
        text.encode("utf-8"),
        text,
        ".md",
    )
    session = get_session()
    material = Material(
        title=req.title,
        category=req.category,
        description=req.description,
        content=req.content,
        content_type=req.content_type,
        tags=req.tags,
        char_count=len(req.content),
        source_file=str(disk_path),
    )
    session.add(material)
    session.commit()
    session.refresh(material)
    return {"id": material.id, "title": material.title, "file_path": str(disk_path), "message": "材料已添加"}


@app.post("/projects/{project_id}/materials/upload")
async def upload_project_material(
    project_id: int,
    category: str = Form(...),
    chapter: Optional[str] = Form(None),
    file: UploadFile = File(...),
):
    """上传项目补充材料，并归档到磁盘材料库供后续匹配使用。"""
    session = get_session()
    project = session.get(Project, project_id)
    if not project:
        raise HTTPException(404, "项目不存在")
    content = await file.read()
    if not content:
        raise HTTPException(400, "上传文件为空")
    text, storage_suffix = _extract_uploaded_material_text(file.filename or "material.md", content)
    disk_path = _write_material_file(category, file.filename or "material.md", content, text, storage_suffix)
    material = Material(
        title=disk_path.stem,
        category=category,
        description=f"项目 {project.name} 上传材料",
        content=text,
        content_type="markdown" if storage_suffix == ".md" else "docx",
        source_file=str(disk_path),
        char_count=len(text),
    )
    session.add(material)
    session.flush()
    active_tender = (
        session.query(Tender)
        .filter(Tender.project_id == project_id, Tender.type == "main")
        .order_by(Tender.id.desc())
        .first()
    )
    usage = MaterialUsage(
        material_id=material.id,
        project_id=project_id,
        tender_id=active_tender.id if active_tender else None,
        chapter=chapter,
    )
    session.add(usage)
    session.commit()
    session.refresh(material)
    return {
        "id": material.id,
        "title": material.title,
        "category": material.category,
        "file_path": str(disk_path),
        "char_count": material.char_count,
        "message": "材料已上传并归档到主标材料库",
    }


# ---- 标书生成（委托给 Orchestrator 自动工作流）----

@app.post("/projects/{project_id}/generate")
def generate_tender(project_id: int, req: GenerateDraftRequest):
    """使用 Orchestrator 自动工作流：匹配 → 生成 → 审查 → (可选)陪标。"""
    orch = Orchestrator(tender_config)
    result = orch.run_workflow(
        project_id=project_id,
        tender_type=req.tender_type,
        need_sub_bid=req.need_sub_bid,
    )

    if "error" in result:
        raise HTTPException(400, result["error"])

    return {
        "message": "标书生成完成",
        "parsed": result.get("parsed"),
        "matches": result.get("matches"),
        "draft": result.get("draft"),
        "main_review": result.get("main_review"),
        "sub_draft": result.get("sub_draft"),
        "sub_review": result.get("sub_review"),
    }


# ---- 标书文件树 / 单文件读取（FileSidebar 用）----

@app.get("/projects/{project_id}/tenders/{tender_id}/files")
def list_tender_files(project_id: int, tender_id: int):
    """返回标书文件树（按 outline 的 6 分类聚合），供右侧 FileSidebar 渲染。"""
    session = get_session()
    project = session.get(Project, project_id)
    tender = session.get(Tender, tender_id)
    if not project or not tender or tender.project_id != project_id:
        raise HTTPException(404, "项目或标书不存在")
    if not tender.draft_path:
        return {"tender_id": tender_id, "type": tender.type,
                "root": "", "folders": [], "top_files": []}

    main_dir = Path(tender.draft_path).parent
    if not main_dir.exists():
        return {"tender_id": tender_id, "type": tender.type,
                "root": str(main_dir), "folders": [], "top_files": []}

    folders: list[dict] = []
    top_files: list[dict] = []
    # 顶层固定文件
    for fname in ("cover.md", "draft.md", "deviation.md"):
        fpath = main_dir / fname
        if fpath.exists():
            top_files.append({
                "name": fname,
                "path": fname,
                "size": fpath.stat().st_size,
                "kind": "top",
            })

    def _folder_entry(folder: Path, prefix: str = "") -> dict:
        rel = f"{prefix}/{folder.name}" if prefix else folder.name
        cat_no = ""
        cat_name = folder.name
        if "_" in folder.name:
            head, _, rest = folder.name.partition("_")
            if head.isdigit():
                cat_no = head
                cat_name = rest
        volume_label = {
            "commercial": "商务文件",
            "technical": "技术文件",
            "price": "商务文件",
            "other": "商务文件",
        }.get(prefix, prefix)
        display_name = f"{volume_label} / {cat_name}" if prefix else cat_name
        files = []
        for f in sorted(folder.iterdir(), key=lambda p: p.name):
            if f.is_file() and f.suffix == ".md":
                fname = f.name
                chap_no = None
                if "_" in fname:
                    head = fname.split("_", 1)[0]
                    if head.isdigit():
                        chap_no = int(head)
                files.append({
                    "name": fname,
                    "path": f"{rel}/{fname}",
                    "size": f.stat().st_size,
                    "chapter_no": chap_no,
                })
        return {
            "name": display_name,
            "category": rel,
            "category_no": cat_no,
            "path": rel,
            "files": files,
        }

    # 分类子目录：兼容旧结构 main/<category>/ 与新结构 main/<volume>/<category>/
    for entry in sorted(main_dir.iterdir(), key=lambda p: p.name):
        if not entry.is_dir():
            continue
        child_dirs = [p for p in entry.iterdir() if p.is_dir()]
        if entry.name in {"commercial", "technical", "price", "other"} and child_dirs:
            for cat_dir in sorted(child_dirs, key=lambda p: p.name):
                folders.append(_folder_entry(cat_dir, prefix=entry.name))
        else:
            folders.append(_folder_entry(entry))

    return {
        "tender_id": tender_id,
        "type": tender.type,
        "root": str(main_dir),
        "folders": folders,
        "top_files": top_files,
    }


@app.get("/projects/{project_id}/tenders/{tender_id}/files/{file_path:path}")
def read_tender_file(project_id: int, tender_id: int, file_path: str):
    """读取单个 .md 文件的原文（Markdown 文本），供全屏查看器渲染。

    路径安全：必须落在 tender.draft_path 父目录内,防止 ../ 越权。
    """
    session = get_session()
    project = session.get(Project, project_id)
    tender = session.get(Tender, tender_id)
    if not project or not tender or tender.project_id != project_id:
        raise HTTPException(404, "项目或标书不存在")
    if not tender.draft_path:
        raise HTTPException(404, "标书未生成文件")

    main_dir = Path(tender.draft_path).parent.resolve()
    target = (main_dir / file_path).resolve()
    try:
        target.relative_to(main_dir)
    except ValueError:
        raise HTTPException(403, "路径越界")

    if not target.exists() or not target.is_file():
        raise HTTPException(404, f"文件不存在: {file_path}")
    if target.suffix.lower() != ".md":
        raise HTTPException(400, "仅支持 .md 文件")

    return PlainTextResponse(
        target.read_text(encoding="utf-8"),
        media_type="text/markdown; charset=utf-8",
    )


@app.put("/projects/{project_id}/tenders/{tender_id}/files/{file_path:path}")
def save_tender_file(project_id: int, tender_id: int, file_path: str, req: MarkdownContentRequest):
    """保存单个生成后的 .md 文件。"""
    session = get_session()
    project = session.get(Project, project_id)
    tender = session.get(Tender, tender_id)
    if not project or not tender or tender.project_id != project_id:
        raise HTTPException(404, "项目或标书不存在")
    if not tender.draft_path:
        raise HTTPException(404, "标书未生成文件")

    main_dir = Path(tender.draft_path).parent.resolve()
    target = (main_dir / file_path).resolve()
    try:
        target.relative_to(main_dir)
    except ValueError:
        raise HTTPException(403, "路径越界")

    if target.suffix.lower() != ".md":
        raise HTTPException(400, "仅支持保存 .md 文件")
    if not target.exists() or not target.is_file():
        raise HTTPException(404, f"文件不存在: {file_path}")

    target.write_text(req.content, encoding="utf-8")
    if target.name == "deviation.md":
        tender.deviation_path = str(target)
    if target.name == "draft.md":
        tender.draft_path = str(target)
    session.commit()
    return {"message": "文件已保存", "path": file_path, "size": len(req.content)}


@app.get("/projects/{project_id}/match")
def match_materials(project_id: int, tender_type: str = "main"):
    """用已确认的提纲做材料匹配（前提：提纲已确认；如未确认则自动生成一份）。"""
    session = get_session()
    project = session.get(Project, project_id)
    if not project:
        raise HTTPException(404, "项目不存在")

    orch = Orchestrator(tender_config)
    orch.ctx.project_id = project_id
    orch.ctx.tender_type = tender_type
    orch.ctx.parsed_data = json.loads(project.parsed_data) if project.parsed_data else {}
    orch.ctx.outline = (
        orch.ctx.parsed_data.get("_confirmed_outline")
        or orch.ctx.parsed_data.get("_generated_outline")
        or []
    )

    # 兜底：没有 outline 就先生成
    if not orch.ctx.outline:
        outline_result = orch.agents["matcher"].generate_outline(orch.ctx)
        orch.ctx.outline = outline_result.get("outline", [])

    match_result = orch.agents["matcher"].match_materials(orch.ctx)
    orch.ctx.parsed_data["_confirmed_outline"] = orch.ctx.outline
    orch.ctx.parsed_data["chapters"] = match_result.get("chapters", [])

    project.parsed_data = json.dumps(orch.ctx.parsed_data, ensure_ascii=False)
    project.status = "materials_preparing"
    session.commit()

    return {
        "message": "材料匹配完成，请确认每章选择：",
        "chapters": match_result.get("chapters", []),
        "action_hint": "可以告诉我需要替换哪个章节，或说'继续'进入生成"
    }


@app.patch("/projects/{project_id}/match/chapter")
def update_matched_chapter(project_id: int, req: MatchChapterUpdateRequest):
    """保存用户手动调整后的章节-材料匹配。"""
    session = get_session()
    project = session.get(Project, project_id)
    if not project:
        raise HTTPException(404, "项目不存在")
    parsed = json.loads(project.parsed_data) if project.parsed_data else {}
    chapters = parsed.get("chapters")
    if not isinstance(chapters, list):
        raise HTTPException(400, "当前项目还没有材料匹配结果")

    matched = None
    for chapter in chapters:
        if isinstance(chapter, dict) and chapter.get("chapter_id") == req.chapter_id:
            matched = chapter
            break
    if matched is None:
        raise HTTPException(404, "章节匹配项不存在")

    if req.file_path:
        path = _material_path_from_request(req.file_path)
        matched["file_path"] = str(path)
        matched["material_title"] = req.material_title or path.stem
    else:
        matched["file_path"] = None
        matched["material_title"] = req.material_title or "无匹配材料"
    matched["match_score"] = req.match_score or "高"
    matched["reason"] = req.reason or "用户手动选择"
    matched["alternatives"] = [
        alt for alt in matched.get("alternatives", [])
        if isinstance(alt, dict) and alt.get("file_path") != matched.get("file_path")
    ]
    matched_sources = matched.get("matched_sources") if isinstance(matched.get("matched_sources"), list) else []
    if req.file_path:
        matched_sources = [
            {
                "source_type": "material_library",
                "title": matched["material_title"],
                "file_path": matched["file_path"],
                "match_score": matched["match_score"],
                "reason": matched["reason"],
                "evidence": [],
            }
        ] + [
            source for source in matched_sources
            if not (isinstance(source, dict) and source.get("source_type") == "material_library")
        ]
    matched["matched_sources"] = matched_sources

    parsed["chapters"] = chapters
    project.parsed_data = json.dumps(parsed, ensure_ascii=False)
    project.status = "materials_preparing"
    session.commit()
    return {"message": "匹配结果已更新", "chapter": matched}


@app.post("/projects/{project_id}/outline/export")
def export_outline(project_id: int, req: ExportRequest):
    """导出提纲为 Markdown（提纲阶段即可使用）。"""
    session = get_session()
    project = session.get(Project, project_id)
    if not project:
        raise HTTPException(404, "项目不存在")
    result = _export_outline_file(project, req.format)
    return {
        "message": f"已导出提纲为 {result['format']}",
        **result,
    }


@app.post("/tenders/{tender_id}/review")
def review_tender(tender_id: int):
    """由 ReviewerAgent 执行 C01-C10 终审检查。"""
    session = get_session()
    tender = session.get(Tender, tender_id)
    if not tender:
        raise HTTPException(404, "标书不存在")

    orch = Orchestrator(tender_config)
    orch.ctx.tender_id = tender_id
    orch.ctx.tender_type = tender.type
    orch.ctx.project_id = tender.project_id

    review_result = orch.agents["reviewer"].execute(orch.ctx)
    checks = review_result.get("checks", [])
    summary = review_result.get("summary", {})
    project = session.get(Project, tender.project_id)
    next_status = "review_failed" if review_result.get("error") else "reviewed"
    if project:
        project.status = next_status
    tender.status = next_status
    session.commit()

    review_error = review_result.get("error")
    return {
        "tender_id": tender_id,
        "checks": checks,
        "issues": review_result.get("issues", []),
        "summary": summary,
        "deterministic_count": review_result.get("deterministic_count", 0),
        "error": review_error,
        "message": (
            f"终审失败：{review_error}"
            if review_error
            else f"终审检查完成，发现 {summary.get('medium', 0)} 个警告项"
        ),
        "action_hint": (
            "请先修复终审前置条件后重试"
            if review_error
            else "要我一键修正，还是你手动处理？"
        )
    }


@app.post("/tenders/{tender_id}/export")
def export_tender(tender_id: int, req: ExportRequest):
    """导出标书为 Markdown / Word / PDF。"""
    session = get_session()
    tender = session.get(Tender, tender_id)
    if not tender:
        raise HTTPException(404, "标书不存在")

    result = _export_tender_file(tender, req.format)

    return {
        "message": f"已导出为 {result['format']}",
        **result,
    }


@app.get("/downloads/outlines/{project_id}/{fmt}")
def download_outline_export(project_id: int, fmt: str):
    export_format = _normalize_export_format(fmt)
    if export_format != "markdown":
        raise HTTPException(400, "提纲导出暂只支持 Markdown")
    path = EXPORTS_DIR / f"outline_{project_id}.md"
    if not path.exists():
        raise HTTPException(404, "提纲导出文件不存在，请先导出")
    return FileResponse(
        path,
        media_type="text/markdown; charset=utf-8",
        filename=path.name,
    )


@app.get("/downloads/{tender_id}/{fmt}")
def download_export(tender_id: int, fmt: str):
    export_format = _normalize_export_format(fmt)
    suffix = "docx" if export_format == "word" else "md"
    path = EXPORTS_DIR / f"tender_{tender_id}.{suffix}"
    if not path.exists():
        raise HTTPException(404, "导出文件不存在，请先导出")
    media_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if export_format == "word"
        else "text/markdown; charset=utf-8"
    )
    return FileResponse(path, media_type=media_type, filename=path.name)


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
